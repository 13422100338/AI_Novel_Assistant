# main_window.py
import os
import docx
import re
from PyQt6.QtWidgets import (QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, QLabel,
                             QTextEdit, QPushButton, QScrollArea, QSplitter, QMessageBox,
                             QFileDialog, QTreeWidget, QTreeWidgetItem, QMenu, QStackedWidget,
                             QInputDialog, QToolBar, QCheckBox, QGroupBox)
from PyQt6.QtCore import Qt, QSettings
from PyQt6.QtGui import QShortcut, QKeySequence, QAction, QTextDocument, QTextCursor
from PyQt6.QtPrintSupport import QPrinter
from data_manager import NovelProject
from ai_worker import AutoPilotWorker, AIWorker, CorrectionWorker, SummaryWorker, SegmentModifyWorker, ChapterAnalysisWorker
from context_builder import NovelContextBuilder
from ui_components import SettingsDialog, CharacterWidget
from ui_helpers import CollapsibleSection, render_chat_messages_html
from model_profiles import get_model_profile
from novel_memory import NovelMemoryStore
from PyQt6.QtWidgets import QToolButton, QMenu, QListWidget, QDockWidget # 新增引用

class MainWindow(QMainWindow):
    def __init__(self, project_path):
        super().__init__()
        self.project = NovelProject(project_path)
        self.memory = NovelMemoryStore(project_path)
        self.settings = QSettings("AIWriter", "Settings")
        self.character_widgets = []
        self.current_vol_index = -1
        self.current_chap_index = -1
        self.switch_project = False
        self.is_generating = False
        self.is_generating_summaries = False  # <--- 新增这行
        self.import_analysis_queue = []
        self.analysis_target = None
        self.last_context_token_estimate = 0
        self.last_output_token_estimate = 0
        self.plot_chat_messages = []

        self.gen_v_idx = -1  # 正在生成的卷索引
        self.gen_c_idx = -1  # 正在生成的章索引
        self.gen_content_buffer = ""  # 正文生成的内存缓冲区
        self.gen_reasoning_buffer = ""  # 思考过程的内存缓冲区

        self.setWindowTitle(f"AI 网文辅助创作系统 - 📖 [{self.project.meta['title']}] (按 Ctrl+S 保存)")
        self.resize(1400, 850)

        self.init_menu_and_toolbar()
        self.init_ui()
        self.setup_shortcuts()
        self.refresh_tree()

    def init_menu_and_toolbar(self):
        # 菜单栏
        menubar = self.menuBar()
        file_menu = menubar.addMenu('文件')

        settings_action = QAction('⚙️ 全局/大模型设置', self)
        settings_action.triggered.connect(self.open_settings)
        file_menu.addAction(settings_action)

        # 显眼的顶部工具栏 (任何时候都可以快速调出设置)
        toolbar = QToolBar("Main Toolbar")
        toolbar.setMovable(False)
        self.addToolBar(toolbar)

        # ====== 返回首页按钮 ======
        btn_home = QPushButton("🏠 返回首页")
        btn_home.setStyleSheet(
            "background-color: transparent; border: 1px solid #DCDFE6; font-weight:bold; color: #E6A23C;")
        btn_home.clicked.connect(self.return_to_home)
        toolbar.addWidget(btn_home)

        # ====== 一键成书按钮 ======
        btn_export = QPushButton("📚 一键成书")
        btn_export.setStyleSheet(
            "background-color: transparent; border: 1px solid #DCDFE6; font-weight:bold; color: #67C23A;")
        btn_export.clicked.connect(self.export_book)
        toolbar.addWidget(btn_export)

        btn_import = QPushButton("📥 导入已有稿件")
        btn_import.setStyleSheet(
            "background-color: transparent; border: 1px solid #DCDFE6; font-weight:bold; color: #409EFF;")
        btn_import.clicked.connect(self.import_manuscript_files)
        toolbar.addWidget(btn_import)

        toolbar.addSeparator()

        btn_settings = QPushButton("⚙️ 设置模型参数")
        btn_settings.setStyleSheet("background-color: transparent; border: 1px solid #DCDFE6; font-weight:bold;")
        btn_settings.clicked.connect(self.open_settings)
        toolbar.addWidget(btn_settings)

        toolbar.addSeparator()

        lbl_status = QLabel("  💡 提示：在左侧树状图右键可新建卷/章。")
        lbl_status.setStyleSheet("color: #909399; font-size: 13px;")
        toolbar.addWidget(lbl_status)

        # ====== 自动挂机按钮 (升级为下拉菜单) ======
        self.btn_auto_pilot = QToolButton()
        self.btn_auto_pilot.setText("🤖 开启自动挂机")
        self.btn_auto_pilot.setPopupMode(QToolButton.ToolButtonPopupMode.InstantPopup)
        self.btn_auto_pilot.setStyleSheet(
            "background-color: transparent; border: 1px solid #DCDFE6; font-weight:bold; color: #9C27B0; padding: 5px;")

        self.auto_pilot_menu = QMenu(self)
        self.auto_pilot_menu.addAction("📚 一键生成全书", lambda: self.toggle_auto_pilot("full"))
        self.auto_pilot_menu.addAction("📄 一键生成本卷", lambda: self.toggle_auto_pilot("volume"))
        self.btn_auto_pilot.setMenu(self.auto_pilot_menu)
        toolbar.addWidget(self.btn_auto_pilot)

        # ====== 全文一键纠错菜单 ======
        self.btn_full_correct = QToolButton()
        self.btn_full_correct.setText("🩺 全文一键纠错")
        self.btn_full_correct.setPopupMode(QToolButton.ToolButtonPopupMode.InstantPopup)
        self.btn_full_correct.setStyleSheet(
            "background-color: transparent; border: 1px solid #DCDFE6; font-weight:bold; color: #F56C6C; padding: 5px;")

        self.full_correct_menu = QMenu(self)
        self.full_correct_menu.addAction("错别字/语病纠错", lambda: self.start_correction("full", "typo"))
        self.full_correct_menu.addAction("设定/逻辑纠错", lambda: self.start_correction("full", "setting"))
        self.full_correct_menu.addAction("🌟 全部纠错", lambda: self.start_correction("full", "all"))
        self.btn_full_correct.setMenu(self.full_correct_menu)

        toolbar.addWidget(self.btn_full_correct)

        # # ====== 侧边栏开关 ======（v2版本换了位置）
        # self.btn_toggle_log = QPushButton("📋 纠错日志")
        # self.btn_toggle_log.setStyleSheet("background-color: transparent; border: none; color: #909399;")
        # self.btn_toggle_log.setCheckable(True)
        # self.btn_toggle_log.clicked.connect(self.toggle_log_sidebar)
        # toolbar.addWidget(self.btn_toggle_log)

    def open_settings(self):
        SettingsDialog(self).exec()

    def init_ui(self):
        central_widget = QWidget()
        central_widget.setStyleSheet("""
            QWidget { background-color: #FFFFFF; color: #1F2937; font-family: "Microsoft YaHei", "Segoe UI", sans-serif; }
            QGroupBox {
                border: 1px solid #E4E7ED;
                border-radius: 10px;
                margin-top: 10px;
                padding: 8px;
                font-weight: 600;
            }
            QGroupBox::title { subcontrol-origin: margin; left: 10px; padding: 0 4px; }
            QTextEdit, QTreeWidget, QListWidget {
                border: 1px solid #E4E7ED;
                border-radius: 12px;
                background-color: #FAFAFA;
                selection-background-color: #BFDBFE;
            }
            QPushButton, QToolButton {
                border-radius: 12px;
                padding: 8px 12px;
                border: 1px solid #D1D5DB;
                background-color: #FFFFFF;
            }
            QPushButton:hover, QToolButton:hover {
                background-color: #F3F4F6;
                border-color: #9CA3AF;
            }
            QSplitter::handle { background-color: #EEF2F7; }
        """)  # 统一第二版界面风格
        self.setCentralWidget(central_widget)
        main_layout = QHBoxLayout(central_widget)
        main_layout.setContentsMargins(10, 10, 10, 10)

        splitter = QSplitter(Qt.Orientation.Horizontal)
        main_layout.addWidget(splitter)

        # ====== 左侧：文件树导航 ======
        tree_container = QWidget()
        tree_outer_layout = QVBoxLayout(tree_container)
        tree_outer_layout.setContentsMargins(0, 0, 0, 0)
        tree_outer_layout.setSpacing(0)

        left_scroll = QScrollArea()
        left_scroll.setWidgetResizable(True)
        left_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        left_scroll.setStyleSheet("QScrollArea { border: none; background: transparent; }")
        left_content = QWidget()
        tree_layout = QVBoxLayout(left_content)
        tree_layout.setContentsMargins(0, 0, 6, 0)
        tree_layout.setSpacing(8)
        left_scroll.setWidget(left_content)
        tree_outer_layout.addWidget(left_scroll)

        self.tree = QTreeWidget()
        self.tree.setHeaderLabel("📚 小说大纲目录 (右键操作)")
        self.tree.header().setStyleSheet("font-weight: bold; font-size: 15px; color: #303133;")
        self.tree.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.tree.customContextMenuRequested.connect(self.show_context_menu)
        self.tree.itemClicked.connect(self.on_tree_select)
        self.tree.setMinimumHeight(260)
        tree_layout.addWidget(self._make_collapsible_section("章节管理", self.tree, True))

        self.current_character_state = QTextEdit()
        self.current_character_state.setReadOnly(True)
        self.current_character_state.setMinimumHeight(180)
        self.current_character_state.setPlaceholderText("选择章节后，这里显示人物在当前节点的心理、动机、目标和最近行动。")
        self.current_character_state.setStyleSheet(
            "background-color: #F8FAFC; border: 1px solid #E4E7ED; border-radius: 8px; padding: 8px; color: #334155;"
        )
        tree_layout.addWidget(self._make_collapsible_section("当前节点人物心理", self.current_character_state, True))

        # ====== 中间：设定面板区 (Stacked) ======
        self.stacked_widget = QStackedWidget()

        # 页面0: 全局设定
        self.page_global = QWidget()
        gl_layout = QVBoxLayout(self.page_global)
        gl_layout.setContentsMargins(10, 0, 10, 0)
        gl_layout.addWidget(QLabel("<span style='font-size:16px; font-weight:bold;'>🌍 全局故事梗概</span>"))
        self.story_synopsis_input = QTextEdit(self.project.meta["global_synopsis"])
        gl_layout.addWidget(self.story_synopsis_input)

        gl_layout.addWidget(
            QLabel("<span style='font-size:16px; font-weight:bold; margin-top:10px;'>👥 核心人物设定</span>"))
        self.char_list_layout = QVBoxLayout()
        scroll_char = QScrollArea()
        scroll_char.setWidgetResizable(True)
        scroll_char.setMinimumHeight(260)
        scroll_char.setStyleSheet("border: none;")
        char_container = QWidget()
        char_container.setLayout(self.char_list_layout)
        scroll_char.setWidget(char_container)
        gl_layout.addWidget(scroll_char)

        btn_add_char = QPushButton("➕ 添加新人物")
        btn_add_char.setStyleSheet("border-style: dashed; background-color: #FAFAFA;")
        btn_add_char.clicked.connect(self.add_character)
        gl_layout.addWidget(btn_add_char)

        btn_save_global = QPushButton("💾 保存全局设定")
        btn_save_global.setStyleSheet("background-color: #409EFF; color: white; font-weight: bold; border: none;")
        btn_save_global.clicked.connect(self.save_global_meta)
        gl_layout.addWidget(btn_save_global)

        # 页面1: 卷设定
        self.page_volume = QWidget()
        vl_layout = QVBoxLayout(self.page_volume)
        self.lbl_vol_title = QLabel("<b>当前卷: </b>")
        self.lbl_vol_title.setStyleSheet("font-size: 16px; color: #303133;")
        vl_layout.addWidget(self.lbl_vol_title)
        self.vol_synopsis_input = QTextEdit()
        self.vol_synopsis_input.setPlaceholderText("本卷的核心主线、剧情走向梗概...")
        vl_layout.addWidget(self.vol_synopsis_input)
        btn_save_vol = QPushButton("💾 保存卷设定")
        btn_save_vol.setStyleSheet("background-color: #409EFF; color: white; font-weight: bold; border: none;")
        btn_save_vol.clicked.connect(self.save_vol_meta)
        vl_layout.addWidget(btn_save_vol)

        # 页面2: 章设定
        self.page_chapter = QWidget()
        cl_layout = QVBoxLayout(self.page_chapter)
        self.lbl_chap_title = QLabel("<b>当前章: </b>")
        self.lbl_chap_title.setStyleSheet("font-size: 16px; color: #303133;")
        cl_layout.addWidget(self.lbl_chap_title)
        self.chap_synopsis_input = QTextEdit()
        self.chap_synopsis_input.setPlaceholderText("本章细纲、出场人物、名场面要求...")
        cl_layout.addWidget(self.chap_synopsis_input)
        btn_save_chap = QPushButton("💾 保存章设定")
        btn_save_chap.setStyleSheet("background-color: #409EFF; color: white; font-weight: bold; border: none;")
        btn_save_chap.clicked.connect(self.save_chap_meta)
        cl_layout.addWidget(btn_save_chap)

        self.stacked_widget.addWidget(self.page_global)
        self.stacked_widget.addWidget(self.page_volume)
        self.stacked_widget.addWidget(self.page_chapter)
        self.stacked_widget.setMinimumHeight(520)
        tree_layout.addWidget(self._make_collapsible_section("设定 / 细纲 / 人物介绍", self.stacked_widget, True))
        tree_layout.addStretch()

        # ====== 右侧：写作输出区 ======
        right_widget = QWidget()
        right_layout = QVBoxLayout(right_widget)
        right_layout.setContentsMargins(10, 0, 0, 0)

        self.plot_chat_box = QGroupBox("剧情聊天")
        self.plot_chat_box.setStyleSheet(
            "QGroupBox { background:#FFFFFF; border:1px solid #E5E7EB; border-radius:18px; "
            "margin-top:10px; padding:12px; font-weight:700; color:#111827; }"
            "QGroupBox::title { subcontrol-origin: margin; left:14px; padding:0 6px; }"
        )
        plot_layout = QVBoxLayout(self.plot_chat_box)
        plot_layout.setContentsMargins(12, 18, 12, 12)
        plot_layout.setSpacing(10)
        self.plot_chat_output = QTextEdit()
        self.plot_chat_output.setReadOnly(True)
        self.plot_chat_output.setAcceptRichText(True)
        self.plot_chat_output.setPlaceholderText("这里用于和 AI 商讨剧情走向、人物动机、伏笔回收和章节节奏。")
        self.plot_chat_output.setMinimumHeight(360)
        self.plot_chat_output.setStyleSheet(
            "QTextEdit { background-color: #FFFFFF; border: none; border-radius: 14px; padding: 8px; }"
        )
        plot_layout.addWidget(self.plot_chat_output)

        plot_input_row = QHBoxLayout()
        plot_input_row.setSpacing(8)
        self.plot_chat_input = QTextEdit()
        self.plot_chat_input.setPlaceholderText("给剧情搭档发消息……")
        self.plot_chat_input.setFixedHeight(74)
        self.plot_chat_input.setStyleSheet(
            "QTextEdit { background:#F9FAFB; border:1px solid #D1D5DB; border-radius:18px; "
            "padding:12px; color:#111827; }"
            "QTextEdit:focus { border:1px solid #111827; background:#FFFFFF; }"
        )
        self.btn_plot_chat = QPushButton("发送")
        self.btn_plot_chat.setMinimumWidth(78)
        self.btn_plot_chat.setStyleSheet(
            "QPushButton { background:#111827; color:white; border:none; border-radius:18px; "
            "font-weight:700; padding:12px 16px; }"
            "QPushButton:hover { background:#374151; }"
            "QPushButton:disabled { background:#CBD5E1; color:#F8FAFC; }"
        )
        self.btn_plot_chat.clicked.connect(self.start_plot_chat)
        plot_input_row.addWidget(self.plot_chat_input)
        plot_input_row.addWidget(self.btn_plot_chat)
        plot_layout.addLayout(plot_input_row)
        chat_widget = QWidget()
        chat_layout = QVBoxLayout(chat_widget)
        chat_layout.setContentsMargins(10, 0, 0, 0)
        chat_layout.addWidget(self.plot_chat_box)
        self.chat_context_hint = QLabel("聊天 AI 会读取当前正文编辑区内容、当前章设定和长篇记忆。")
        self.chat_context_hint.setStyleSheet("color: #64748B; font-size: 12px;")
        chat_layout.addWidget(self.chat_context_hint)
        self._render_plot_chat()

        # 1. 顶部操作按钮行 (放在同一个水平布局里)
        btn_action_layout = QHBoxLayout()

        self.btn_start = QPushButton("🚀 根据设定撰写本章")
        self.btn_start.setEnabled(False)  # 必须选中章节才能写
        self.btn_start.setStyleSheet(
            "font-size: 16px; font-weight: bold; background-color: #A0CFFF; color: white; border: none; padding: 12px; border-radius: 16px;"
        )
        self.btn_start.clicked.connect(self.start_generation)
        btn_action_layout.addWidget(self.btn_start)

        # 章节一键纠错按钮
        self.btn_chap_correct = QToolButton()
        self.btn_chap_correct.setText("🔧 章节一键纠错")
        self.btn_chap_correct.setPopupMode(QToolButton.ToolButtonPopupMode.InstantPopup)
        self.btn_chap_correct.setStyleSheet(
            "font-size: 14px; font-weight: bold; background-color: #E6A23C; color: white; border: none; padding: 10px; border-radius: 16px;")
        self.btn_chap_correct.setEnabled(False)  # 初始化时未选章节不可用

        self.chap_correct_menu = QMenu(self)
        self.chap_correct_menu.addAction("错别字/语病纠错", lambda: self.start_correction("chapter", "typo"))
        self.chap_correct_menu.addAction("设定/逻辑纠错", lambda: self.start_correction("chapter", "setting"))
        self.chap_correct_menu.addAction("🌟 全部纠错", lambda: self.start_correction("chapter", "all"))
        self.btn_chap_correct.setMenu(self.chap_correct_menu)
        btn_action_layout.addWidget(self.btn_chap_correct)

        # 将按钮行加入右侧主布局
        right_layout.addLayout(btn_action_layout)

        info_row = QHBoxLayout()
        self.btn_plot_points = QPushButton("🧭 情节点")
        self.btn_plot_points.setToolTip("展示 AI 为当前章节整理的关键情节点、伏笔和章节摘要。")
        self.btn_plot_points.setStyleSheet(
            "QPushButton { background:#EEF2FF; color:#3730A3; border:1px solid #C7D2FE; "
            "border-radius:14px; font-weight:700; padding:8px 14px; }"
            "QPushButton:hover { background:#E0E7FF; }"
        )
        self.btn_plot_points.clicked.connect(self.show_current_plot_points)
        info_row.addWidget(self.btn_plot_points)
        self.lbl_token_usage = QLabel("Token 估算：上下文 0 / 输出 0")
        self.lbl_token_usage.setStyleSheet(
            "color: #64748B; background-color: #F8FAFC; border: 1px solid #E4E7ED; border-radius: 16px; padding: 6px;"
        )
        info_row.addWidget(self.lbl_token_usage, 1)
        right_layout.addLayout(info_row)

        # 2. 思考过程显示区
        self.btn_toggle_thinking = QPushButton("🔽 收起思考过程")
        self.btn_toggle_thinking.setStyleSheet(
            "background-color: transparent; border: none; color: #909399; text-align: left;")
        self.btn_toggle_thinking.clicked.connect(self.toggle_thinking)
        right_layout.addWidget(self.btn_toggle_thinking)

        self.thinking_output = QTextEdit()
        self.thinking_output.setReadOnly(True)
        self.thinking_output.setStyleSheet(
            "background-color: #F8F9FA; color: #8A8F99; border: 1px solid #E4E7ED; border-radius: 16px;")
        self.thinking_output.setFixedHeight(120)
        right_layout.addWidget(self.thinking_output)

        # 3. 正文显示区
        right_layout.addWidget(
            QLabel("<span style='font-size:16px; font-weight:bold;'>✍️ 小说正文区 (按 Ctrl+S 实时保存到 docx)</span>"))

        self.content_output = QTextEdit()
        self.content_output.setStyleSheet("""
                    font-size: 16px; 
                    line-height: 1.8; 
                    padding: 15px; 
                    color: #2C3E50;
                    background-color: #FAFAFA;
                """)
        right_layout.addWidget(self.content_output)

        # ====== 【新增】日志侧边栏 (作为一个可隐藏的 QListWidget) ======
        self.log_list = QListWidget()
        self.log_list.setStyleSheet(
            "background-color: #FAFAFA; border: 1px solid #E4E7ED; color: #606266; padding: 5px;")
        self.log_list.setWordWrap(True)  # 【新增】开启自动换行，防止日志过长难以查看
        self.log_list.hide()  # 默认隐藏

        # ====== 给正文输出区绑定右键菜单 ======
        self.content_output.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.content_output.customContextMenuRequested.connect(self.show_editor_context_menu)

        # ====== 【重构】最右侧的竖向工具栏与 Copilot 面板 ======
        self.sidebar_stacked = QStackedWidget()
        self.sidebar_stacked.hide()  # 默认隐藏面板

        # -- 侧边栏 Page 0: 纠错日志 --
        self.log_list = QListWidget()
        self.log_list.setStyleSheet(
            "background-color: #FAFAFA; border: 1px solid #E4E7ED; color: #606266; padding: 5px;")
        self.log_list.setWordWrap(True)
        self.sidebar_stacked.addWidget(self.log_list)

        # -- 侧边栏 Page 1: 文段修正 (Copilot) --
        self.modifier_widget = QWidget()
        mod_layout = QVBoxLayout(self.modifier_widget)
        mod_layout.setContentsMargins(5, 5, 5, 5)
        mod_layout.addWidget(QLabel("<b>🎯 选中片段</b>"))
        self.mod_selected_text = QTextEdit()
        self.mod_selected_text.setReadOnly(True)
        self.mod_selected_text.setFixedHeight(100)
        self.mod_selected_text.setStyleSheet("background-color: #FDFDFD; color: #909399;")
        mod_layout.addWidget(self.mod_selected_text)

        mod_layout.addWidget(QLabel("<b>💬 修改指令 (Prompt)</b>"))
        self.mod_instruction = QTextEdit()
        self.mod_instruction.setPlaceholderText("例如：让语气更暴躁一点、把这段扩写细致一些、换成古风描写...")
        self.mod_instruction.setFixedHeight(80)
        mod_layout.addWidget(self.mod_instruction)

        btn_mod_layout = QHBoxLayout()
        self.btn_submit_modify = QPushButton("✨ 生成修改")
        self.btn_submit_modify.setStyleSheet(
            "background-color: #409EFF; color: white; font-weight: bold; padding: 6px;")
        self.btn_submit_modify.clicked.connect(self.start_segment_modification)
        self.btn_cancel_modify = QPushButton("🛑 停止")
        self.btn_cancel_modify.setStyleSheet("background-color: #F56C6C; color: white; padding: 6px;")
        self.btn_cancel_modify.setEnabled(False)
        self.btn_cancel_modify.clicked.connect(self.cancel_segment_modification)
        btn_mod_layout.addWidget(self.btn_submit_modify)
        btn_mod_layout.addWidget(self.btn_cancel_modify)
        mod_layout.addLayout(btn_mod_layout)

        mod_layout.addWidget(QLabel("<b>🤖 AI 修正结果</b>"))
        self.mod_result = QTextEdit()
        self.mod_result.setReadOnly(True)
        mod_layout.addWidget(self.mod_result)

        self.btn_apply_replace = QPushButton("✅ 替换原文选中片段")
        self.btn_apply_replace.setStyleSheet(
            "background-color: #67C23A; color: white; font-weight: bold; padding: 8px;")
        self.btn_apply_replace.setEnabled(False)
        self.btn_apply_replace.clicked.connect(self.apply_modification)
        mod_layout.addWidget(self.btn_apply_replace)
        self.sidebar_stacked.addWidget(self.modifier_widget)

        # -- 侧边栏 Page 2: 长篇记忆状态 --
        self.memory_widget = QWidget()
        memory_layout = QVBoxLayout(self.memory_widget)
        memory_layout.setContentsMargins(5, 5, 5, 5)
        memory_layout.addWidget(QLabel("<b>长篇记忆库</b>"))
        self.memory_dashboard = QTextEdit()
        self.memory_dashboard.setReadOnly(True)
        self.memory_dashboard.setPlainText(self.memory.dashboard_text())
        memory_layout.addWidget(self.memory_dashboard)
        self.btn_refresh_memory = QPushButton("刷新记忆状态")
        self.btn_refresh_memory.clicked.connect(self.refresh_memory_dashboard)
        self.btn_reanalyze_chapter = QPushButton("重新分析当前章")
        self.btn_reanalyze_chapter.clicked.connect(self.start_reanalyze_current_chapter)
        memory_layout.addWidget(self.btn_refresh_memory)
        memory_layout.addWidget(self.btn_reanalyze_chapter)
        self.sidebar_stacked.addWidget(self.memory_widget)

        # -- 最最右侧的竖向按钮柱 (侧边导航栏) --
        vertical_toolbar = QWidget()
        vertical_toolbar.setFixedWidth(46)  # 稍微放宽一点点，避免文字贴边

        # 专属独立样式：白色背景、去默认边框、选中时左侧显示蓝色指示条
        vertical_toolbar.setStyleSheet("""
                    QWidget {
                        background-color: #FFFFFF;
                        border-left: 1px solid #E4E7ED;
                    }
                    QPushButton {
                        background-color: transparent;
                        border: none;
                        border-radius: 10px;
                        color: #909399;
                        padding: 10px 0;
                        font-size: 13px;
                        line-height: 1.5;
                    }
                    QPushButton:hover {
                        background-color: #F2F6FC;
                        color: #409EFF;
                    }
                    QPushButton:checked {
                        background-color: #ECF5FF;
                        color: #409EFF;
                        font-weight: bold;
                        border-left: 3px solid #409EFF; 
                        border-radius: 0px;
                    }
                """)
        v_toolbar_layout = QVBoxLayout(vertical_toolbar)
        v_toolbar_layout.setContentsMargins(2, 10, 2, 0)  # 顶部留一点空隙
        v_toolbar_layout.setSpacing(10)

        self.btn_sidebar_log = QPushButton("📋\n日\n志")
        self.btn_sidebar_log.setCheckable(True)
        self.btn_sidebar_log.setFixedSize(40, 80)
        self.btn_sidebar_log.clicked.connect(lambda: self.toggle_right_sidebar(0, self.btn_sidebar_log))

        self.btn_sidebar_modifier = QPushButton("🪄\n修\n正")
        self.btn_sidebar_modifier.setCheckable(True)
        self.btn_sidebar_modifier.setFixedSize(40, 80)
        self.btn_sidebar_modifier.clicked.connect(lambda: self.toggle_right_sidebar(1, self.btn_sidebar_modifier))

        self.btn_sidebar_memory = QPushButton("记忆\n库")
        self.btn_sidebar_memory.setCheckable(True)
        self.btn_sidebar_memory.setFixedSize(40, 80)
        self.btn_sidebar_memory.clicked.connect(lambda: self.toggle_right_sidebar(2, self.btn_sidebar_memory))

        v_toolbar_layout.addWidget(self.btn_sidebar_log)
        v_toolbar_layout.addWidget(self.btn_sidebar_modifier)
        v_toolbar_layout.addWidget(self.btn_sidebar_memory)
        v_toolbar_layout.addStretch()  # 把按钮顶在上面

        # 重新拼装主视窗
        right_splitter = QSplitter(Qt.Orientation.Horizontal)
        right_splitter.addWidget(right_widget)  # 写作主力区
        right_splitter.addWidget(self.sidebar_stacked)  # Copilot 浮动面板区
        right_splitter.addWidget(vertical_toolbar)  # 最右侧细条工具栏
        right_splitter.setSizes([750, 250, 45])
        right_splitter.setCollapsible(2, False)  # 不允许收起细长工具栏

        splitter.addWidget(tree_container)
        splitter.addWidget(right_splitter)
        splitter.addWidget(chat_widget)
        splitter.setSizes([400, 760, 420])
        splitter.setStretchFactor(0, 0)
        splitter.setStretchFactor(1, 1)
        splitter.setStretchFactor(2, 0)

        # 初始化加载人物
        for char_data in self.project.meta.get("characters", []):
            self.add_character(char_data)

    def setup_shortcuts(self):
        shortcut_save = QShortcut(QKeySequence("Ctrl+S"), self)
        shortcut_save.activated.connect(self.save_all)
        shortcut_delete = QShortcut(QKeySequence("Delete"), self.tree)
        shortcut_delete.activated.connect(lambda: self.ui_delete_item(self.tree.currentItem()))

    def _make_collapsible_section(self, title, content_widget, expanded=True):
        return CollapsibleSection(title, content_widget, expanded)

    def _render_plot_chat(self):
        if not hasattr(self, "plot_chat_output"):
            return
        self.plot_chat_output.setHtml(render_chat_messages_html(self.plot_chat_messages))
        bar = self.plot_chat_output.verticalScrollBar()
        bar.setValue(bar.maximum())

    def start_plot_chat(self):
        question = self.plot_chat_input.toPlainText().strip()
        if not question:
            return
        profile = get_model_profile("chat", self.settings)
        if not profile.api_key:
            QMessageBox.warning(self, "缺少 API Key", "请先在设置中配置“剧情商讨模型”的 API Key。")
            self.open_settings()
            return

        self.save_all()
        context = self._current_memory_context()
        editor_content = self.content_output.toPlainText().strip()
        if len(editor_content) > 8000:
            editor_content = "……（当前正文较长，前文省略）……\n" + editor_content[-8000:]
        system_prompt = (
            "你是长篇小说剧情策划搭档，负责和作者讨论剧情走向、人物动机、伏笔回收、节奏与冲突。"
            "你只提供创作建议，不直接替作者改写整章正文，除非作者明确要求。"
        )
        user_prompt = f"""
【小说全局设定】
{self.project.meta.get('global_synopsis', '')}

【当前章节与长篇记忆】
{context}

【当前正文编辑区内容】
{editor_content if editor_content else "当前正文编辑区为空。"}

        【作者问题】
        {question}
        """
        self.update_token_usage(context_text=system_prompt + user_prompt)
        self.plot_chat_messages.append({"role": "user", "content": question})
        self.plot_chat_messages.append({"role": "assistant", "content": ""})
        self._render_plot_chat()
        self.plot_chat_input.clear()
        self.btn_plot_chat.setEnabled(False)
        self.plot_worker = AIWorker(
            api_key=profile.api_key,
            base_url=profile.base_url,
            model=profile.model,
            temperature=profile.temperature,
            max_tokens=profile.max_tokens,
            system_prompt=system_prompt,
            user_prompt=user_prompt,
        )
        self.plot_worker.content_signal.connect(self.append_plot_chat)
        self.plot_worker.error_signal.connect(self.handle_error)
        self.plot_worker.finished_signal.connect(self.plot_chat_finished)
        self.plot_worker.start()

    def append_plot_chat(self, text):
        if not self.plot_chat_messages or self.plot_chat_messages[-1].get("role") != "assistant":
            self.plot_chat_messages.append({"role": "assistant", "content": ""})
        self.plot_chat_messages[-1]["content"] += text
        self._render_plot_chat()

    def plot_chat_finished(self):
        self.btn_plot_chat.setEnabled(True)
        self._render_plot_chat()

    def refresh_memory_dashboard(self):
        self.memory_dashboard.setPlainText(self.memory.dashboard_text())

    def _current_memory_context(self):
        if self.current_vol_index == -1:
            return self.memory.build_context("", "")
        vol = self.project.meta["volumes"][self.current_vol_index]
        chap_name = ""
        if self.current_chap_index != -1:
            chap_name = vol["chapters"][self.current_chap_index]["name"]
        return self.memory.build_context(vol["name"], chap_name)

    def refresh_current_character_state(self):
        if not hasattr(self, "current_character_state"):
            return
        if self.current_vol_index == -1:
            self.current_character_state.setPlainText(self.memory.character_state_text())
            return
        vol = self.project.meta["volumes"][self.current_vol_index]
        chap_name = ""
        if self.current_chap_index != -1:
            chap_name = vol["chapters"][self.current_chap_index]["name"]
        self.current_character_state.setPlainText(self.memory.character_state_text(vol["name"], chap_name))

    def show_current_plot_points(self):
        if self.current_vol_index == -1 or self.current_chap_index == -1:
            QMessageBox.information(self, "情节点", "请先选择一个章节。")
            return
        vol = self.project.meta["volumes"][self.current_vol_index]
        chap = vol["chapters"][self.current_chap_index]
        text = self.memory.get_plot_points_text(vol["name"], chap["name"])
        if not text:
            QMessageBox.information(
                self,
                "情节点",
                "当前章节还没有 AI 情节点记录。\n\n你可以点击“重新分析当前章”，或在导入/生成章节后等待自动分析完成。",
            )
            return
        QMessageBox.information(self, f"情节点 - {chap['name']}", text)

    def estimate_tokens(self, text: str) -> int:
        if not text:
            return 0
        ascii_count = sum(1 for ch in text if ord(ch) < 128)
        non_ascii_count = len(text) - ascii_count
        return max(1, int(ascii_count / 4 + non_ascii_count / 1.7))

    def update_token_usage(self, context_text: str = "", output_text: str = ""):
        if context_text:
            self.last_context_token_estimate = self.estimate_tokens(context_text)
        if output_text:
            self.last_output_token_estimate = self.estimate_tokens(output_text)
        if hasattr(self, "lbl_token_usage"):
            self.lbl_token_usage.setText(
                f"Token 估算：上下文 {self.last_context_token_estimate:,} / 输出 {self.last_output_token_estimate:,}"
            )

    def _start_analysis_for_chapter(self, v_idx, c_idx, status_prefix="正在分析章节记忆"):
        profile = get_model_profile("draft", self.settings)
        if not profile.api_key:
            QMessageBox.warning(self, "缺少 API Key", "请先在设置中配置“正文创作模型”的 API Key。")
            self.open_settings()
            return False
        vol = self.project.meta["volumes"][v_idx]
        chap = vol["chapters"][c_idx]
        content = self.project.read_chapter_content(vol["name"], chap["name"])
        if len(content.strip()) < 50:
            return False
        if hasattr(self, "analysis_worker") and self.analysis_worker.isRunning():
            return False
        self.analysis_target = (v_idx, c_idx)
        if hasattr(self, "btn_reanalyze_chapter"):
            self.btn_reanalyze_chapter.setEnabled(False)
        self.statusBar().showMessage(f"{status_prefix}：{vol['name']} / {chap['name']}")
        self.analysis_worker = ChapterAnalysisWorker(
            profile.api_key,
            profile.base_url,
            profile.model,
            min(profile.temperature, 0.8),
            vol["name"],
            chap["name"],
            content,
            self.memory.build_context(vol["name"], chap["name"]),
        )
        self.analysis_worker.status_signal.connect(lambda msg: self.statusBar().showMessage(msg))
        self.analysis_worker.analysis_ready_signal.connect(self.apply_chapter_analysis)
        self.analysis_worker.error_signal.connect(self.handle_error)
        self.analysis_worker.finished_signal.connect(self.analysis_finished)
        self.analysis_worker.start()
        return True

    def start_reanalyze_current_chapter(self):
        if self.current_vol_index == -1 or self.current_chap_index == -1:
            QMessageBox.warning(self, "请选择章节", "请先在左侧选择一个已经有正文的章节。")
            return
        self.save_all()
        vol = self.project.meta["volumes"][self.current_vol_index]
        chap = vol["chapters"][self.current_chap_index]
        content = self.project.read_chapter_content(vol["name"], chap["name"])
        if len(content.strip()) < 50:
            QMessageBox.warning(self, "正文太短", "当前章正文太短，暂时没有足够内容可分析。")
            return
        self._start_analysis_for_chapter(self.current_vol_index, self.current_chap_index, "正在重新分析当前章")
        return

        profile = get_model_profile("draft", self.settings)
        if not profile.api_key:
            QMessageBox.warning(self, "缺少 API Key", "请先在设置中配置“正文创作模型”的 API Key。")
            self.open_settings()
            return

        self.btn_reanalyze_chapter.setEnabled(False)
        self.statusBar().showMessage("正在重新分析当前章...")
        self.analysis_worker = ChapterAnalysisWorker(
            profile.api_key,
            profile.base_url,
            profile.model,
            min(profile.temperature, 0.8),
            vol["name"],
            chap["name"],
            content,
            self._current_memory_context(),
        )
        self.analysis_worker.status_signal.connect(lambda msg: self.statusBar().showMessage(msg))
        self.analysis_worker.analysis_ready_signal.connect(self.apply_chapter_analysis)
        self.analysis_worker.error_signal.connect(self.handle_error)
        self.analysis_worker.finished_signal.connect(self.analysis_finished)
        self.analysis_worker.start()

    def apply_chapter_analysis(self, payload):
        if self.analysis_target:
            v_idx, c_idx = self.analysis_target
        else:
            v_idx, c_idx = self.current_vol_index, self.current_chap_index
        if v_idx == -1 or c_idx == -1:
            return
        vol = self.project.meta["volumes"][v_idx]
        chap = vol["chapters"][c_idx]
        self.memory.upsert_analysis(vol["name"], chap["name"], payload)
        summary = payload.get("summary", "")
        if summary:
            chap["ai_synopsis"] = summary
            self.project.save_meta()
            if v_idx == self.current_vol_index and c_idx == self.current_chap_index:
                self.chap_synopsis_input.setPlainText(chap.get("synopsis", ""))
        self.refresh_memory_dashboard()
        self.refresh_current_character_state()

    def analysis_finished(self):
        self.btn_reanalyze_chapter.setEnabled(True)
        self.analysis_target = None
        if self.import_analysis_queue:
            next_v_idx, next_c_idx = self.import_analysis_queue[0]
            if self._start_analysis_for_chapter(next_v_idx, next_c_idx, "正在分析导入稿件"):
                self.import_analysis_queue.pop(0)
                return
        self.statusBar().showMessage("章节记忆分析完成。", 5000)

    # --- UI 辅助与交互逻辑 ---
    def add_character(self, init_data=None):
        widget = CharacterWidget(self.remove_character, init_data)
        self.char_list_layout.addWidget(widget)
        self.character_widgets.append(widget)

    def remove_character(self, widget):
        self.char_list_layout.removeWidget(widget)
        widget.deleteLater()
        self.character_widgets.remove(widget)

    def toggle_thinking(self):
        is_visible = self.thinking_output.isVisible()
        self.thinking_output.setVisible(not is_visible)
        self.btn_toggle_thinking.setText("🔽 收起思考过程" if not is_visible else "▶️ 展开思考过程")

    # === 日志侧边栏切换 ===（v2版本）
    def toggle_log_sidebar(self, checked):
        if checked:
            self.sidebar_stacked.setCurrentIndex(0)
            self.sidebar_stacked.show()
            self.btn_sidebar_log.setChecked(True)
            self.btn_sidebar_modifier.setChecked(False)
        else:
            self.sidebar_stacked.hide()
            self.btn_sidebar_log.setChecked(False)

    def update_ui_state(self):
        # 1. 检查当前视角的章节是否正在被大模型撰写、纠错或正在自动挂机
        is_auto_piloting = getattr(self, 'is_auto_piloting', False)
        is_correcting = getattr(self, 'is_correcting', False)
        is_generating = getattr(self, 'is_generating', False)
        is_generating_summaries = getattr(self, 'is_generating_summaries', False)  # 取出新状态

        is_viewing_gen_chapter = ((is_generating or is_auto_piloting) and
                                  self.current_vol_index == getattr(self, 'gen_v_idx', -1) and
                                  self.current_chap_index == getattr(self, 'gen_c_idx', -1))

        # 只要正在生成当前章，或者处于全局挂机、全文/单章纠错状态，严格锁定文本框为只读
        self.content_output.setReadOnly(is_viewing_gen_chapter or is_auto_piloting or is_correcting)

        # ====== 状态覆盖优先级判定 ======
        # ====== 状态覆盖优先级判定 ======
        if is_generating_summaries:
            self.btn_start.setEnabled(True)
            self.btn_start.setText("🛑 停止补全总结")
            self.btn_start.setStyleSheet(
                "font-size: 15px; font-weight: bold; background-color: #E6A23C; color: white; border: none; padding: 12px; border-radius: 16px;")

            self.btn_auto_pilot.setEnabled(True)
            self.btn_auto_pilot.setText("🛑 停止补全总结")
            self.btn_auto_pilot.setMenu(None)  # 隐藏菜单
            self.btn_auto_pilot.setStyleSheet(
                "background-color: #E6A23C; border: 1px solid #DCDFE6; font-weight:bold; color: white; padding: 5px;")
            try:
                self.btn_auto_pilot.clicked.disconnect()
            except Exception:
                pass
            self.btn_auto_pilot.clicked.connect(lambda: self.toggle_auto_pilot("stop"))

            self.btn_full_correct.setEnabled(False)
            self.btn_chap_correct.setEnabled(False)

        # 2. 动态改变生成按钮的颜色和文案
        elif is_generating:
            # (保留您原本的 is_generating 逻辑)
            self.btn_start.setEnabled(True)
            if is_viewing_gen_chapter:
                self.btn_start.setText("🛑 停止生成 (正在输出当前章)")
                self.btn_start.setStyleSheet(
                    "font-size: 15px; font-weight: bold; background-color: #F56C6C; color: white; border: none; padding: 12px; border-radius: 16px;")
            else:
                self.btn_start.setText("🛑 停止后台生成 (其他章正在码字)")
                self.btn_start.setStyleSheet(
                    "font-size: 15px; font-weight: bold; background-color: #E6A23C; color: white; border: none; padding: 12px; border-radius: 16px;")

        elif is_auto_piloting:
            self.btn_start.setEnabled(False)
            self.btn_start.setText("🤖 挂机模式进行中...")
            self.btn_start.setStyleSheet(
                "font-size: 16px; font-weight: bold; background-color: #A0CFFF; color: white; border: none; padding: 12px; border-radius: 16px;")

            # 把挂机按钮变成红色停止按钮
            self.btn_auto_pilot.setEnabled(True)
            self.btn_auto_pilot.setText("🛑 停止挂机")
            self.btn_auto_pilot.setMenu(None)
            self.btn_auto_pilot.setStyleSheet(
                "background-color: #F56C6C; border: 1px solid #DCDFE6; font-weight:bold; color: white; padding: 5px;")
            try:
                self.btn_auto_pilot.clicked.disconnect()
            except Exception:
                pass
            self.btn_auto_pilot.clicked.connect(lambda: self.toggle_auto_pilot("stop"))
        else:
            if self.current_chap_index != -1:
                self.btn_start.setText("🚀 根据设定撰写本章")
                self.btn_start.setEnabled(True)
                self.btn_start.setStyleSheet(
                    "font-size: 16px; font-weight: bold; background-color: #67C23A; color: white; border: none; padding: 12px; border-radius: 16px;")
            else:
                self.btn_start.setText("🚀 根据设定撰写本章")
                self.btn_start.setEnabled(False)
                self.btn_start.setStyleSheet(
                    "font-size: 16px; font-weight: bold; background-color: #A0CFFF; color: white; border: none; padding: 12px; border-radius: 16px;")

            # 恢复挂机按钮
            self.btn_auto_pilot.setEnabled(True)
            self.btn_auto_pilot.setText("🤖 开启自动挂机")
            self.btn_auto_pilot.setMenu(self.auto_pilot_menu)
            self.btn_auto_pilot.setStyleSheet(
                "background-color: transparent; border: 1px solid #DCDFE6; font-weight:bold; color: #9C27B0; padding: 5px;")
            try:
                self.btn_auto_pilot.clicked.disconnect()
            except Exception:
                pass

        # ====== 纠错按钮状态更新 ====== (保留下半部分的这一行，不做修改)
        has_chap_selected = (self.current_chap_index != -1)
        is_correcting = getattr(self, 'is_correcting', False)

        if is_correcting:
            # 纠错进行中：把全文纠错按钮魔改成红色的“停止”按钮
            self.btn_full_correct.setEnabled(True)
            self.btn_full_correct.setText("🛑 停止纠错")
            self.btn_full_correct.setMenu(None)  # 摘掉下拉菜单，变成普通可点击按钮
            self.btn_full_correct.setStyleSheet(
                "background-color: #F56C6C; border: 1px solid #DCDFE6; font-weight:bold; color: white; padding: 5px;"
            )
            # 绑定停止事件（先静默解绑以防重复绑定）
            try:
                self.btn_full_correct.clicked.disconnect()
            except Exception:
                pass
            self.btn_full_correct.clicked.connect(self.cancel_correction)

            # 把另一个纠错按钮彻底锁死
            self.btn_chap_correct.setEnabled(False)
            self.btn_chap_correct.setText("🔧 纠错运行中...")
        else:
            # 恢复正常：把菜单装回去，改回原来的颜色
            self.btn_full_correct.setEnabled(True)
            self.btn_full_correct.setText("🩺 全文一键纠错")
            self.btn_full_correct.setMenu(self.full_correct_menu)
            self.btn_full_correct.setStyleSheet(
                "background-color: transparent; border: 1px solid #DCDFE6; font-weight:bold; color: #F56C6C; padding: 5px;"
            )
            # 解除点击事件的绑定，恢复菜单展开默认行为
            try:
                self.btn_full_correct.clicked.disconnect()
            except Exception:
                pass

            self.btn_chap_correct.setEnabled(has_chap_selected)
            self.btn_chap_correct.setText("🔧 章节一键纠错")
            self.btn_chap_correct.setMenu(self.chap_correct_menu)
            self.btn_chap_correct.setStyleSheet(
                "font-size: 14px; font-weight: bold; background-color: #E6A23C; color: white; border: none; padding: 10px; border-radius: 16px;"
            )

    def start_correction(self, scope, mode):
        profile = get_model_profile("draft", self.settings)
        if not profile.api_key:
            QMessageBox.warning(self, "错误", "缺少 API Key！")
            return

        if scope == "chapter" and (self.current_vol_index == -1 or self.current_chap_index == -1):
            QMessageBox.warning(self, "错误", "请先在左侧选择要纠错的章节！")
            return

        warning_msg = "全文纠错将扫描全书摘要，可能消耗大量 Token 且耗时较长，请耐心等待。确认开始吗？" if scope == "full" else "即将使用 AI 检查并修改当前章节原文，确认开始吗？"
        if QMessageBox.question(self, '启动纠错', warning_msg) != QMessageBox.StandardButton.Yes:
            return

        self.save_all()  # 强制保存当前最新状态
        self.is_correcting = True
        self.update_ui_state()

        # 自动展开并清空侧边栏准备记录
        self.btn_toggle_log.setChecked(True)
        self.toggle_log_sidebar(True)
        self.log_list.addItem(f"=== 开始新的纠错任务 ({'全书' if scope == 'full' else '单章'}) ===")
        self.log_list.scrollToBottom()

        base_url = self.settings.value("base_url", "https://api.deepseek.com")
        model = self.settings.value("model", "deepseek-reasoner")
        temp = float(self.settings.value("temperature", 0.7))

        self.correct_worker = CorrectionWorker(api_key, base_url, model, temp, self.project, scope, mode)
        if scope == "chapter":
            self.correct_worker.set_target(self.current_vol_index, self.current_chap_index)

        self.correct_worker.status_signal.connect(lambda msg: self.statusBar().showMessage(msg))
        self.correct_worker.log_signal.connect(self.append_correction_log)
        self.correct_worker.update_text_signal.connect(self.apply_corrected_text)
        self.correct_worker.finished_signal.connect(self.correction_finished)
        self.correct_worker.error_signal.connect(self.handle_error)

        # ====== 【新增】连接思考过程信号，并在启动时清空且展开思考面板 ======
        self.correct_worker.reasoning_signal.connect(self.append_thinking)
        self.thinking_output.clear()
        if not self.thinking_output.isVisible():
            self.toggle_thinking()

        self.correct_worker.start()

    def append_correction_log(self, log_msg):
        self.log_list.addItem(log_msg)
        self.log_list.scrollToBottom()

    def apply_corrected_text(self, v_idx, c_idx, new_content, new_summary):
        vol_name = self.project.meta["volumes"][v_idx]["name"]
        chap = self.project.meta["volumes"][v_idx]["chapters"][c_idx]
        chap_name = chap["name"]

        # 后台落盘
        self.project.save_chapter_content(vol_name, chap_name, new_content)
        if new_summary and new_summary != chap.get("ai_synopsis", ""):
            chap["ai_synopsis"] = new_summary
            self.project.save_meta()

        # 如果当前 UI 正好停留在被修改的这一章，实时刷新文本框
        if self.current_vol_index == v_idx and self.current_chap_index == c_idx:
            self.content_output.setText(new_content)
            self.statusBar().showMessage(f"✨ 当前章节 [{chap_name}] 纠错并刷新完毕！", 3000)

    def correction_finished(self):
        self.is_correcting = False
        self.update_ui_state()  # 这一步会让按钮从“停止”重新变回“一键纠错”

        # 判断是被手动停止的还是自然跑完的
        if hasattr(self, 'correct_worker') and self.correct_worker._is_cancelled:
            self.statusBar().showMessage("🛑 纠错任务已手动终止！", 3000)
            self.log_list.addItem("=== 纠错已手动终止 ===")
        else:
            self.statusBar().showMessage("✅ 纠错任务全部完成！", 3000)
            self.log_list.addItem("=== 纠错任务结束 ===")

        self.log_list.scrollToBottom()

    # --- 目录树逻辑 ---
    def refresh_tree(self):
        self.tree.clear()
        root = QTreeWidgetItem(self.tree, [self.project.meta["title"]])
        root.setIcon(0, self.style().standardIcon(self.style().StandardPixmap.SP_DirIcon))
        root.setData(0, Qt.ItemDataRole.UserRole, {"type": "root"})

        for v_idx, vol in enumerate(self.project.meta["volumes"]):
            v_node = QTreeWidgetItem(root, [vol["name"]])
            v_node.setIcon(0, self.style().standardIcon(self.style().StandardPixmap.SP_FileDialogDetailedView))
            v_node.setData(0, Qt.ItemDataRole.UserRole, {"type": "volume", "v_idx": v_idx})

            for c_idx, chap in enumerate(vol["chapters"]):
                c_node = QTreeWidgetItem(v_node, [chap["name"]])
                c_node.setIcon(0, self.style().standardIcon(self.style().StandardPixmap.SP_FileIcon))
                c_node.setData(0, Qt.ItemDataRole.UserRole, {"type": "chapter", "v_idx": v_idx, "c_idx": c_idx})
        self.tree.expandAll()

    def show_context_menu(self, position):
        item = self.tree.itemAt(position)
        menu = QMenu()
        menu.setStyleSheet(
            "QMenu { background-color: white; border: 1px solid #DCDFE6; } QMenu::item:selected { background-color: #ECF5FF; color: #409EFF; }")

        if not item or item.data(0, Qt.ItemDataRole.UserRole)["type"] == "root":
            action_add_vol = menu.addAction("📁 新建卷")
            action_add_vol.triggered.connect(self.ui_add_volume)
        elif item.data(0, Qt.ItemDataRole.UserRole)["type"] == "volume":
            action_add_chap = menu.addAction("📄 在此卷下新建章")
            v_idx = item.data(0, Qt.ItemDataRole.UserRole)["v_idx"]
            action_add_chap.triggered.connect(lambda: self.ui_add_chapter(v_idx))
            # 【新增】卷的修改与删除
            action_rename = menu.addAction("✏️ 重命名卷")
            action_rename.triggered.connect(lambda: self.ui_rename_item(item))
            action_delete = menu.addAction("🗑️ 删除卷")
            action_delete.triggered.connect(lambda: self.ui_delete_item(item))

        elif item.data(0, Qt.ItemDataRole.UserRole)["type"] == "chapter":
            # 【新增】章的修改与删除
            action_rename = menu.addAction("✏️ 重命名章")
            action_rename.triggered.connect(lambda: self.ui_rename_item(item))
            action_delete = menu.addAction("🗑️ 删除章")
            action_delete.triggered.connect(lambda: self.ui_delete_item(item))

        menu.exec(self.tree.viewport().mapToGlobal(position))

    def ui_add_volume(self):
        text, ok = QInputDialog.getText(self, "新建卷", "请输入卷名:")
        if ok and text:
            self.project.add_volume(text)
            self.refresh_tree()

    def ui_add_chapter(self, v_idx):
        text, ok = QInputDialog.getText(self, "新建章", "请输入章名:")
        if ok and text:
            self.project.add_chapter(v_idx, text)
            self.refresh_tree()

    def ui_rename_item(self, item):
        if not item: return
        data = item.data(0, Qt.ItemDataRole.UserRole)
        if data["type"] == "root": return

        old_name = item.text(0)
        item_type = "卷" if data["type"] == "volume" else "章"

        new_name, ok = QInputDialog.getText(self, f"重命名{item_type}", f"请输入新的{item_type}名:", text=old_name)
        if ok and new_name and new_name.strip() != old_name:
            new_name = new_name.strip()
            # 执行数据重命名
            if data["type"] == "volume":
                self.project.rename_volume(data["v_idx"], new_name)
            elif data["type"] == "chapter":
                self.project.rename_chapter(data["v_idx"], data["c_idx"], new_name)

            # 刷新树与右侧标题显示
            self.refresh_tree()
            if data["type"] == "volume" and self.current_vol_index == data["v_idx"]:
                self.lbl_vol_title.setText(f"<b>当前卷: {new_name}</b>")
            elif data["type"] == "chapter" and self.current_vol_index == data["v_idx"] and self.current_chap_index == \
                    data["c_idx"]:
                vol_name = self.project.meta["volumes"][data["v_idx"]]["name"]
                self.lbl_chap_title.setText(f"<b>当前章: {vol_name} - {new_name}</b>")

    def ui_delete_item(self, item):
        if not item: return
        data = item.data(0, Qt.ItemDataRole.UserRole)
        if data["type"] == "root": return

        if self.is_generating:
            if (data["type"] == "volume" and data["v_idx"] == self.gen_v_idx) or \
                    (data["type"] == "chapter" and data["v_idx"] == self.gen_v_idx and data["c_idx"] == self.gen_c_idx):
                QMessageBox.warning(self, "操作受限", "该卷/章正在后台疯狂码字中，请先停止生成后再尝试删除！")
                return

        item_type = "卷" if data["type"] == "volume" else "章"
        item_name = item.text(0)

        # 读取用户是否开启了“删除前确认”设置
        needs_confirm = self.settings.value("confirm_delete", True, type=bool)

        if needs_confirm:
            msgBox = QMessageBox(self)
            msgBox.setWindowTitle("确认删除")
            msgBox.setText(f"确定要删除{item_type}【{item_name}】吗？\n删除操作同时会移除本地文件，且不可恢复！")
            msgBox.setIcon(QMessageBox.Icon.Warning)
            msgBox.setStandardButtons(QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
            msgBox.setDefaultButton(QMessageBox.StandardButton.No)

            # 植入“不再提醒”的 CheckBox
            cb = QCheckBox("以后不再提醒")
            msgBox.setCheckBox(cb)

            if msgBox.exec() != QMessageBox.StandardButton.Yes:
                return  # 用户取消了删除

            # 如果勾选了不再提醒，更新全局设置
            if cb.isChecked():
                self.settings.setValue("confirm_delete", False)

        # 执行删除
        if data["type"] == "volume":
            self.project.delete_volume(data["v_idx"])
        elif data["type"] == "chapter":
            self.project.delete_chapter(data["v_idx"], data["c_idx"])

        # 删除后，重置右侧编辑面板回到全局设定页
        self.stacked_widget.setCurrentIndex(0)
        self.current_vol_index = -1
        self.current_chap_index = -1
        self.update_ui_state()
        self.refresh_tree()

    def on_tree_select(self, item):
        # 【新增】在切换目录之前，先静默保存当前选中的卷/章信息，防止内容丢失
        self.save_all(silent=True)

        data = item.data(0, Qt.ItemDataRole.UserRole)
        self.current_vol_index = -1
        self.current_chap_index = -1

        if data["type"] == "root":
            self.stacked_widget.setCurrentIndex(0)

        elif data["type"] == "volume":
            v_idx = data["v_idx"]
            self.current_vol_index = v_idx
            vol_data = self.project.meta["volumes"][v_idx]
            self.lbl_vol_title.setText(f"<b>当前卷: {vol_data['name']}</b>")
            self.vol_synopsis_input.setText(vol_data.get("synopsis", ""))
            self.stacked_widget.setCurrentIndex(1)

        elif data["type"] == "chapter":
            v_idx = data["v_idx"]
            c_idx = data["c_idx"]
            self.current_vol_index = v_idx
            self.current_chap_index = c_idx

            vol_data = self.project.meta["volumes"][v_idx]
            chap_data = vol_data["chapters"][c_idx]

            self.lbl_chap_title.setText(f"<b>当前章: {vol_data['name']} - {chap_data['name']}</b>")
            self.chap_synopsis_input.setText(chap_data.get("synopsis", ""))
            self.stacked_widget.setCurrentIndex(2)

            # 【关键修复】：增加对挂机状态 is_auto_piloting 的判断，否则挂机时会被当作普通查看，导致思考过程被 clear()
            is_active_gen = getattr(self, 'is_generating', False) or getattr(self, 'is_auto_piloting', False)

            if is_active_gen and getattr(self, 'gen_v_idx', -1) == v_idx and getattr(self, 'gen_c_idx', -1) == c_idx:
                # 如果切回了正在生成的章，展示内存中的实时流
                self.content_output.setText(self.gen_content_buffer)
                self.thinking_output.setText(self.gen_reasoning_buffer)
                # 滚动条移到最底端
                self.content_output.moveCursor(self.content_output.textCursor().MoveOperation.End)
                self.thinking_output.moveCursor(self.thinking_output.textCursor().MoveOperation.End)
            else:
                # 查看其他章节，读取本地记录
                content = self.project.read_chapter_content(vol_data["name"], chap_data["name"])
                self.content_output.setText(content)
                self.thinking_output.clear()
        self.refresh_current_character_state()
        if hasattr(self, "content_output"):
            self.update_token_usage(output_text=self.content_output.toPlainText())
        self.update_ui_state()

    def return_to_home(self):
        reply = QMessageBox.question(self, '返回首页', '确定要退出当前项目并返回首页吗？\n(系统将自动保存当前进度)',
                                     QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
                                     QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.Yes:
            self.save_all()  # 自动保存当前数据
            self.switch_project = True  # 设置标志位为 True
            self.close()  # 关闭当前主窗口

    # --- 数据保存逻辑 ---
    def save_global_meta(self, silent=False):
        self.project.meta["global_synopsis"] = self.story_synopsis_input.toPlainText().strip()
        chars = []
        for w in self.character_widgets:
            d = w.get_data()
            if any(d.values()):
                chars.append(d)
        self.project.meta["characters"] = chars
        self.project.save_meta()
        if not silent:
            QMessageBox.information(self, "提示", "全局设定保存成功！")

    def save_vol_meta(self, silent=False):
        if self.current_vol_index != -1:
            self.project.meta["volumes"][self.current_vol_index][
                "synopsis"] = self.vol_synopsis_input.toPlainText().strip()
            self.project.save_meta()
            if not silent:
                QMessageBox.information(self, "提示", "当前卷设定保存成功！")

    def save_chap_meta(self, silent=False):
        if self.current_chap_index != -1:
            self.project.meta["volumes"][self.current_vol_index]["chapters"][self.current_chap_index][
                "synopsis"] = self.chap_synopsis_input.toPlainText().strip()
            self.project.save_meta()
            if not silent:
                QMessageBox.information(self, "提示", "当前章设定保存成功！")

    def save_all(self, silent=True):
        # 1. 如果在全局页，保存全局；如果在卷页，保存卷；如果在章页，保存章梗概和正文
        idx = self.stacked_widget.currentIndex()
        if idx == 0:
            self.save_global_meta(silent=silent)
        elif idx == 1:
            self.save_vol_meta(silent=silent)
        elif idx == 2:
            self.save_chap_meta(silent=silent)
            # 保存 docx 正文
            if self.current_vol_index != -1 and self.current_chap_index != -1:
                vol_name = self.project.meta["volumes"][self.current_vol_index]["name"]
                chap_name = self.project.meta["volumes"][self.current_vol_index]["chapters"][self.current_chap_index][
                    "name"]
                new_content = self.content_output.toPlainText()
                old_content = self.project.read_chapter_content(vol_name, chap_name)
                if old_content.strip() and old_content != new_content:
                    self.memory.archive_chapter_version(vol_name, chap_name, old_content, "manual-save")
                self.project.save_chapter_content(vol_name, chap_name, new_content)

        # 无论是快捷键还是切换章节触发，都在底部状态栏提供无感提示
        self.statusBar().showMessage("✅ 小说正文及设定已自动保存！", 3000)

    def import_manuscript_files(self):
        file_paths, _ = QFileDialog.getOpenFileNames(
            self,
            "导入已有稿件",
            "",
            "Manuscript Files (*.txt *.docx);;Text Files (*.txt);;Word Files (*.docx)"
        )
        if not file_paths:
            return

        self.save_all()
        target_v_idx = self._ensure_import_volume()
        imported = []
        for path in file_paths:
            try:
                text = self._read_import_file(path)
            except Exception as e:
                QMessageBox.warning(self, "导入失败", f"无法读取文件：\n{path}\n\n{e}")
                continue
            chapters = self._split_imported_text(text, os.path.splitext(os.path.basename(path))[0])
            for title, content in chapters:
                chap_name = self._unique_chapter_name(target_v_idx, title)
                synopsis = "由导入稿件自动创建，等待 AI 生成压缩摘要。"
                self.project.add_chapter(target_v_idx, chap_name, synopsis=synopsis, ai_synopsis="")
                c_idx = len(self.project.meta["volumes"][target_v_idx]["chapters"]) - 1
                self.project.save_chapter_content(self.project.meta["volumes"][target_v_idx]["name"], chap_name, content)
                imported.append((target_v_idx, c_idx))

        if not imported:
            return
        self.refresh_tree()
        self.import_analysis_queue.extend(imported)
        QMessageBox.information(
            self,
            "导入完成",
            f"已导入 {len(imported)} 个章节。\n\n系统会自动提取上下文压缩摘要、情节点、人物存档和伏笔账本。"
        )
        if not (hasattr(self, "analysis_worker") and self.analysis_worker.isRunning()):
            next_v_idx, next_c_idx = self.import_analysis_queue[0]
            if self._start_analysis_for_chapter(next_v_idx, next_c_idx, "正在分析导入稿件"):
                self.import_analysis_queue.pop(0)

    def _ensure_import_volume(self):
        for idx, vol in enumerate(self.project.meta["volumes"]):
            if vol["name"] == "导入稿件":
                return idx
        self.project.add_volume("导入稿件", "从 TXT/DOCX 导入的既有稿件。")
        return len(self.project.meta["volumes"]) - 1

    def _read_import_file(self, path):
        ext = os.path.splitext(path)[1].lower()
        if ext == ".docx":
            document = docx.Document(path)
            return "\n".join(p.text for p in document.paragraphs if p.text.strip())
        for encoding in ("utf-8", "utf-8-sig", "gbk"):
            try:
                with open(path, "r", encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()

    def _split_imported_text(self, text, fallback_title):
        text = text.replace("\r\n", "\n").replace("\r", "\n").strip()
        if not text:
            return []
        pattern = re.compile(r"(?m)^(第[零一二三四五六七八九十百千万\d]+[章节卷回][^\n]{0,40}|Chapter\s+\d+[^\n]{0,40})\s*$")
        matches = list(pattern.finditer(text))
        chapters = []
        if matches:
            for idx, match in enumerate(matches):
                title = match.group(1).strip()
                start = match.end()
                end = matches[idx + 1].start() if idx + 1 < len(matches) else len(text)
                content = text[start:end].strip()
                if content:
                    chapters.append((title, content))
        if chapters:
            return chapters
        return [(fallback_title or "导入章节", text)]

    def _unique_chapter_name(self, v_idx, base_name):
        safe = re.sub(r'[\\/:*?"<>|]', "_", (base_name or "导入章节").strip())[:60] or "导入章节"
        existing = {c["name"] for c in self.project.meta["volumes"][v_idx]["chapters"]}
        if safe not in existing:
            return safe
        counter = 2
        while f"{safe}-{counter}" in existing:
            counter += 1
        return f"{safe}-{counter}"

    def export_book(self):
        # 1. 强制保存当前最新进度
        self.save_all()

        # 2. 弹出保存文件对话框
        file_path, filter_type = QFileDialog.getSaveFileName(
            self,
            "一键成书 - 选择导出位置",
            f"{self.project.meta['title']}.docx",
            "Word 文档 (*.docx);;Markdown 文档 (*.md);;纯文本 (*.txt);;PDF 文档 (*.pdf)"
        )

        if not file_path:
            return

        # 3. 根据后缀名调用相应的导出方法
        try:
            ext = os.path.splitext(file_path)[1].lower()
            title = self.project.meta['title']

            if ext == '.docx':
                self._export_docx(file_path, title)
            elif ext == '.md':
                self._export_md(file_path, title)
            elif ext == '.txt':
                self._export_txt(file_path, title)
            elif ext == '.pdf':
                self._export_pdf(file_path, title)

            QMessageBox.information(self, "导出成功", f"恭喜！小说已成功导出至：\n{file_path}")
        except Exception as e:
            QMessageBox.critical(self, "导出失败", f"导出过程中发生错误：\n{str(e)}")

    def _export_docx(self, file_path, title):
        doc = docx.Document()
        doc.add_heading(title, 0)  # 书名作为主标题

        for vol in self.project.meta["volumes"]:
            doc.add_heading(vol["name"], level=1)  # 卷名作为一级标题
            for chap in vol["chapters"]:
                doc.add_heading(chap["name"], level=2)  # 章名作为二级标题
                content = self.project.read_chapter_content(vol["name"], chap["name"])
                for line in content.split('\n'):
                    if line.strip():
                        doc.add_paragraph(line.strip())
        doc.save(file_path)

    def _export_txt(self, file_path, title):
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(f"《{title}》\n\n")
            for vol in self.project.meta["volumes"]:
                f.write(f"【{vol['name']}】\n\n")
                for chap in vol["chapters"]:
                    f.write(f"  {chap['name']}\n\n")
                    content = self.project.read_chapter_content(vol["name"], chap["name"])
                    f.write(f"{content}\n\n")

    def _export_md(self, file_path, title):
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(f"# {title}\n\n")
            for vol in self.project.meta["volumes"]:
                f.write(f"## {vol['name']}\n\n")
                for chap in vol["chapters"]:
                    f.write(f"### {chap['name']}\n\n")
                    content = self.project.read_chapter_content(vol["name"], chap["name"])
                    f.write(f"{content}\n\n")

    def _export_pdf(self, file_path, title):
        # PDF 导出利用 PyQt6 自带的富文本转换为 HTML 再渲染打印的机制
        html_content = f"<h1 style='text-align: center;'>{title}</h1>"
        for vol in self.project.meta["volumes"]:
            html_content += f"<h2 style='color: #2C3E50;'>{vol['name']}</h2>"
            for chap in vol["chapters"]:
                html_content += f"<h3>{chap['name']}</h3>"
                content = self.project.read_chapter_content(vol["name"], chap["name"])
                for line in content.split('\n'):
                    if line.strip():
                        html_content += f"<p style='text-indent: 2em; line-height: 1.5;'>{line.strip()}</p>"

        document = QTextDocument()
        document.setHtml(html_content)

        printer = QPrinter(QPrinter.PrinterMode.HighResolution)
        printer.setOutputFormat(QPrinter.OutputFormat.PdfFormat)
        printer.setOutputFileName(file_path)

        # 渲染生成 PDF
        document.print(printer)

    def _chapter_position_list(self):
        positions = []
        for v_idx, vol in enumerate(self.project.meta["volumes"]):
            for c_idx, chap in enumerate(vol["chapters"]):
                positions.append((v_idx, c_idx, vol["name"], chap["name"]))
        return positions

    def _recent_full_chapters_context(self, v_idx, c_idx, count=3):
        positions = self._chapter_position_list()
        try:
            current_pos = next(i for i, item in enumerate(positions) if item[0] == v_idx and item[1] == c_idx)
        except StopIteration:
            return ""
        recent = positions[max(0, current_pos - count):current_pos]
        blocks = []
        for _, _, vol_name, chap_name in recent:
            content = self.project.read_chapter_content(vol_name, chap_name).strip()
            if not content:
                continue
            # 内部安全阀：保留最近章节的正文风格，但避免单章异常超长拖垮请求。
            if len(content) > 12000:
                content = content[-12000:]
            blocks.append(f"【{vol_name} / {chap_name}】\n{content}")
        return "\n\n".join(blocks)

    def _compressed_history_before_recent(self, v_idx, c_idx, recent_count=3):
        positions = self._chapter_position_list()
        try:
            current_pos = next(i for i, item in enumerate(positions) if item[0] == v_idx and item[1] == c_idx)
        except StopIteration:
            return ""
        cutoff = max(0, current_pos - recent_count)
        lines = []
        for old_v_idx, old_c_idx, vol_name, chap_name in positions[:cutoff]:
            chap = self.project.meta["volumes"][old_v_idx]["chapters"][old_c_idx]
            summary = chap.get("ai_synopsis") or chap.get("synopsis") or ""
            if summary.strip():
                lines.append(f"- {vol_name} / {chap_name}: {summary.strip()}")
        return "\n".join(lines)

    # --- 核心大模型生成逻辑 (包含复杂的上下文组装) ---
    def build_prompts(self):
        return NovelContextBuilder(self.project, self.memory).build_prompts(
            self.current_vol_index,
            self.current_chap_index,
        )

        meta = self.project.meta

        # 1. 组装全局设定
        global_story = meta.get("global_synopsis", "未提供。")
        char_texts = [f"【{c['name']}】 性别:{c['gender']} 性格:{c['personality']} 经历:{c['experience']}" for c in
                      meta.get("characters", [])]
        char_setting = "\n".join(char_texts) if char_texts else "未提供明确人物。"

        system_prompt = f"""你是一位经验丰富的网文大神作家。请根据全局设定和上下文连贯地撰写小说正文。

            【全局故事大纲】
            {global_story}

            【核心人物设定】
            {char_setting}

            【写作要求】
            1. 严格遵循世界观、人设和剧情逻辑。
            2. 动作、神态、心理描写生动，符合网文爽感节奏。
            3. 【强制指令】字数必须严格限定在2500-3500字之间！爽点必须密集，严禁任何无用的废话和无效表达。
            4. 直接输出正文，禁止任何解释性废话和多余的寒暄。
            5.对话要口语化，多用短句，讲话方式符合人设，拒绝‘翻译腔’。角色说话要有情绪和潜台词，不要像写说明书或做思想汇报一样客观中立。特别注意：不要出现‘我无权评价’、‘这取决于你’这类典型的 AI 废话，或者用正常人类不会使用的比喻句等。
            6. 【重要】在正文输出完毕后，必须另起一行并严格以 `[AI_SUMMARY]` 作为分割符，然后输出约500字的本章详细梗概。此部分仅用于系统内部记录。"""

        # 2. 组装历史上下文与上一章内容
        past_context = ""
        prev_chapter_content = ""

        v_idx = self.current_vol_index
        c_idx = self.current_chap_index

        prev_v_idx, prev_c_idx = -1, -1
        if c_idx > 0:
            prev_v_idx, prev_c_idx = v_idx, c_idx - 1
        elif v_idx > 0:
            for i in range(v_idx - 1, -1, -1):
                if len(meta["volumes"][i]["chapters"]) > 0:
                    prev_v_idx = i
                    prev_c_idx = len(meta["volumes"][i]["chapters"]) - 1
                    break

        if prev_v_idx != -1 and prev_c_idx != -1:
            pv_name = meta["volumes"][prev_v_idx]["name"]
            pc_name = meta["volumes"][prev_v_idx]["chapters"][prev_c_idx]["name"]
            prev_chapter_content = self.project.read_chapter_content(pv_name, pc_name)
            if len(prev_chapter_content) > 1500:
                prev_chapter_content = "...(前文省略)...\n" + prev_chapter_content[-1500:]

        # 【修改处】提取过往所有梗概时，优先使用 ai_synopsis
        history_str = ""
        for i in range(v_idx + 1):
            vol = meta["volumes"][i]
            history_str += f"\n> {vol['name']} (梗概: {vol.get('synopsis', '无')})\n"

            chap_limit = c_idx if i == v_idx else len(vol["chapters"])
            for j in range(chap_limit):
                chap = vol["chapters"][j]

                # 优先读取 AI 之前生成的梗概，如果没有则降级读取用户的细纲
                ai_syn = chap.get("ai_synopsis", "")
                user_syn = chap.get("synopsis", "无")
                display_syn = ai_syn if ai_syn.strip() else user_syn

                history_str += f"  - {chap['name']}: {display_syn}\n"

        if len(history_str) > 15000:
            history_str = "【注意：因前文过长，此处仅提供过往卷梗概】\n"
            for i in range(v_idx + 1):
                vol = meta["volumes"][i]
                history_str += f"\n> {vol['name']} (梗概: {vol.get('synopsis', '无')})\n"

        if not history_str.strip():
            history_str = "本书刚刚开篇，无过往历史。"

        curr_vol = meta["volumes"][v_idx]
        curr_chap = curr_vol["chapters"][c_idx]
        memory_context = self.memory.build_context(curr_vol["name"], curr_chap["name"])
        recent_full_context = self._recent_full_chapters_context(v_idx, c_idx, count=3)
        older_compressed_context = self._compressed_history_before_recent(v_idx, c_idx, recent_count=3)

        user_prompt = f"""请为我撰写最新章节的正文。

【较早章节压缩轨迹】
{older_compressed_context.strip() if older_compressed_context.strip() else history_str.strip()}

【最近章节全文，用于保持行文风格、节奏和人物声音】
{recent_full_context if recent_full_context.strip() else "暂无最近章节全文。"}

【长篇压缩记忆 / 人物状态 / 伏笔正典】
{memory_context if memory_context.strip() else "暂无结构化长篇记忆。"}

"""
        if prev_chapter_content.strip():
            user_prompt += f"""【本次写作任务】
            当前所处卷：{curr_vol['name']}
            本卷核心梗概：{curr_vol.get('synopsis', '无')}

            当前需撰写章节：{curr_chap['name']}
            本章细纲要求：{curr_chap.get('synopsis', '无')}

            【行动指令】
            请根据本章细纲要求，顺着上一章的情节展开。
            严格限定正文长度在2000-3000字之间，确保爽点密集、拒绝水文，扩写为文笔流畅的完整正文！
            【重要】在正文输出完毕后，必须另起一行并严格以 `[AI_SUMMARY]` 作为分割符，然后输出约500字高度结构化的【本章复盘与记忆锚点】
                在 `[AI_SUMMARY]` 之后，必须严格按照以下3个维度输出：
                1. 核心剧情脉络：按时间顺序简述本章发生的实质性事件。
                2. 人物状态更新：记录本章主角及配角的行为及心态。
                3. 物品设定更新：记录本章所有物品状态
            """
        else:
            user_prompt += f"""【本次写作任务】
            当前所处卷：{curr_vol['name']}
            本卷核心梗概：{curr_vol.get('synopsis', '无')}

            当前需撰写章节：{curr_chap['name']}
            本章细纲要求：{curr_chap.get('synopsis', '无')}

            【行动指令】
            请根据全局设定、本卷梗概和本章细纲展开。若这是开篇，请迅速建立人物处境、核心冲突和可持续推进的悬念。
            严格限定正文长度在2000-3000字之间，确保爽点密集、拒绝水文，扩写为文笔流畅的完整正文！
            【重要】在正文输出完毕后，必须另起一行并严格以 `[AI_SUMMARY]` 作为分割符，然后输出约500字高度结构化的【本章复盘与记忆锚点】。
            """

        return system_prompt, user_prompt

    def start_generation(self):
        # 1. 拦截正在进行总结补全时的取消操作
        if getattr(self, 'is_generating_summaries', False):
            if hasattr(self, 'summary_worker') and self.summary_worker.isRunning():
                self.summary_worker.cancel()
            self.btn_start.setText("🛑 正在停止补全...")
            self.btn_start.setEnabled(False)
            return

        # 2. 拦截正在进行普通生成时的取消操作
        if getattr(self, 'is_generating', False):
            if hasattr(self, 'worker') and self.worker.isRunning():
                self.worker.cancel()
            self.btn_start.setText("🛑 正在停止...")
            self.btn_start.setEnabled(False)
            return

        api_key = self.settings.value("api_key", "")
        if not api_key:
            QMessageBox.warning(self, "错误", "缺少 API Key，请点击上方【⚙️ 设置模型参数】按钮进行配置！")
            self.open_settings()
            return

        # 强制保存后，进入检查流水线。
        # 这里传入目标章节：只会检查排在它“前面”的内容
        self.save_all()
        self._check_and_fill_summaries(self.current_vol_index, self.current_chap_index, self._execute_start_generation)

    def _execute_start_generation(self):
        """真正的原单章挂机逻辑"""
        system_prompt, user_prompt = self.build_prompts()
        self.update_token_usage(context_text=system_prompt + user_prompt, output_text="")

        self.is_generating = True
        self.gen_v_idx = self.current_vol_index
        self.gen_c_idx = self.current_chap_index
        self.gen_content_buffer = ""
        self.gen_reasoning_buffer = ""

        self.content_output.clear()
        self.thinking_output.clear()

        self.hit_summary_delimiter = False
        self.update_ui_state()

        profile = get_model_profile("draft", self.settings)
        self.worker = AIWorker(api_key=profile.api_key, base_url=profile.base_url, model=profile.model,
                               temperature=profile.temperature, max_tokens=profile.max_tokens, system_prompt=system_prompt,
                               user_prompt=user_prompt)
        self.worker.reasoning_signal.connect(self.append_thinking)
        self.worker.content_signal.connect(self.append_content)
        self.worker.error_signal.connect(self.handle_error)
        self.worker.finished_signal.connect(self.generation_finished)
        self.worker.start()

    def _check_and_fill_summaries(self, target_v_idx, target_c_idx, callback):
        """
        核心拦截器：检查前面所有章节是否有缺失的AI总结。如果有，先启动 SummaryWorker。
        target_v_idx, target_c_idx: 目标章节。如果是自动挂机，传 None, None，即检查全书所有已有内容的章节。
        callback: 补全完成后要接着调用的原生方法（_execute_start_generation 或 _execute_auto_pilot）
        """
        tasks = []
        for v_idx, vol in enumerate(self.project.meta["volumes"]):
            for c_idx, chap in enumerate(vol["chapters"]):
                # 如果是单章生成，只需要检查目标章节“之前”的章节
                if target_v_idx is not None and target_c_idx is not None:
                    if v_idx > target_v_idx or (v_idx == target_v_idx and c_idx >= target_c_idx):
                        continue  # 跳过目标章节本身及之后的所有章

                ai_syn = chap.get("ai_synopsis", "").strip()
                if not ai_syn:
                    # 读取本地看是否真的写了正文
                    content = self.project.read_chapter_content(vol["name"], chap["name"])
                    if len(content.strip()) > 100:  # 正文>100字才算有内容需要总结
                        tasks.append({
                            "v_idx": v_idx, "c_idx": c_idx,
                            "vol_name": vol["name"], "chap_name": chap["name"],
                            "content": content
                        })

        if not tasks:
            # 没有任何缺失，直接执行原本的任务
            callback()
            return

        # --- 发现缺失，启动补全工作流 ---
        self.is_generating_summaries = True
        self.update_ui_state()
        self.statusBar().showMessage(f"⏳ 发现 {len(tasks)} 个已写章节缺失 AI 总结，正在自动补全以免影响记忆...")

        profile = get_model_profile("draft", self.settings)
        self.summary_worker = SummaryWorker(profile.api_key, profile.base_url, profile.model, profile.temperature, tasks)
        self.summary_worker.status_signal.connect(lambda msg: self.statusBar().showMessage(msg))
        self.summary_worker.summary_ready_signal.connect(self._on_missing_summary_ready)
        self.summary_worker.finished_signal.connect(lambda: self._on_missing_summary_finished(callback))
        self.summary_worker.error_signal.connect(self._on_summary_error)
        self.summary_worker.start()

    def _on_missing_summary_ready(self, v_idx, c_idx, summary):
        # 回写数据到结构中
        chap = self.project.meta["volumes"][v_idx]["chapters"][c_idx]
        chap["ai_synopsis"] = summary
        self.project.save_meta()
        vol_name = self.project.meta["volumes"][v_idx]["name"]
        self.memory.upsert_analysis(vol_name, chap["name"], {"summary": summary})
        self.refresh_memory_dashboard()

        # 顺手把 UI 里可能看得到的界面同步一下（如果用户正停留在该章）
        if self.current_vol_index == v_idx and self.current_chap_index == c_idx:
            if not self.chap_synopsis_input.toPlainText().strip():
                self.chap_synopsis_input.setText(summary)

    def _on_missing_summary_finished(self, callback):
        self.is_generating_summaries = False

        # 如果是被手动停止的，不继续执行回调
        if getattr(self, "summary_worker", None) and self.summary_worker._is_cancelled:
            self.statusBar().showMessage("🛑 总结补全已被手动终止！", 3000)
            self.update_ui_state()
            return

        self.statusBar().showMessage("✅ 缺失的 AI 总结全部补全完毕！", 3000)
        self.update_ui_state()
        callback()  # 触发真正的正文生成逻辑

    def _on_summary_error(self, err_msg):
        self.is_generating_summaries = False
        self.update_ui_state()
        QMessageBox.critical(self, "补全错误", f"补全缺失的章节总结时发生异常，已终止后续操作：\n{err_msg}")

    def append_thinking(self, text):
        self.gen_reasoning_buffer += text  # 永远写进后台缓冲区

        # 【关键修复】：如果是挂机模式且还没开始写具体某章正文(gen_v_idx == -1)，说明在全局规划，强制展示思考过程
        is_planning = getattr(self, 'is_auto_piloting', False) and self.gen_v_idx == -1

        # 只有当处于大纲规划期，或用户正停留在正在生成的具体章节时，才实时渲染在屏幕上
        if is_planning or (self.current_vol_index == self.gen_v_idx and self.current_chap_index == self.gen_c_idx):
            self.thinking_output.insertPlainText(text)
            self.thinking_output.ensureCursorVisible()

    def append_content(self, text):
        # [cite_start]【关键修复】：实时将 AI 吐出的文字拼接到后台缓冲区中 [cite: 196]
        self.gen_content_buffer += text
        self.update_token_usage(output_text=self.gen_content_buffer.split("[AI_SUMMARY]")[0])

        if "[AI_SUMMARY]" in self.gen_content_buffer:
            if not getattr(self, 'hit_summary_delimiter', False):
                self.hit_summary_delimiter = True
                # 触发分割符时，将正文的最后一部分清理干净渲染到UI上，之后停止更新UI的正文部分
                if self.current_vol_index == self.gen_v_idx and self.current_chap_index == self.gen_c_idx:
                    main_content = self.gen_content_buffer.split("[AI_SUMMARY]")[0].strip()
                    self.content_output.setPlainText(main_content)
                    self.content_output.moveCursor(self.content_output.textCursor().MoveOperation.End)
        else:
            # 正常渲染正文
            if self.current_vol_index == self.gen_v_idx and self.current_chap_index == self.gen_c_idx:
                # 【核心修复】：先强制将光标移动到文本最末尾，再插入文本。防止鼠标乱点导致文字插错位置！
                cursor = self.content_output.textCursor()
                cursor.movePosition(cursor.MoveOperation.End)
                self.content_output.setTextCursor(cursor)

                # 在末尾插入最新文本流
                self.content_output.insertPlainText(text)
                self.content_output.ensureCursorVisible()

    def handle_error(self, err_msg):
        QMessageBox.critical(self, "生成错误", f"请求发生异常：\n{err_msg}")
        # 【修改处】根据当前的模式，调用对应的结束/重置方法
        if getattr(self, 'is_auto_piloting', False):
            self.auto_pilot_finished()
        else:
            self.generation_finished()

    def generation_finished(self):
        if self.gen_v_idx != -1 and self.gen_c_idx != -1:
            vol_name = self.project.meta["volumes"][self.gen_v_idx]["name"]
            chap_data = self.project.meta["volumes"][self.gen_v_idx]["chapters"][self.gen_c_idx]
            chap_name = chap_data["name"]

            # 【核心修改】将缓冲区的内容根据标识符一分为二
            parts = self.gen_content_buffer.split("[AI_SUMMARY]")
            main_content = parts[0].strip()
            ai_summary = parts[1].strip() if len(parts) > 1 else ""

            # 1. 保存纯净的正文到 docx
            old_content = self.project.read_chapter_content(vol_name, chap_name)
            if old_content.strip() and old_content != main_content:
                self.memory.archive_chapter_version(vol_name, chap_name, old_content, "ai-generation")
            self.project.save_chapter_content(vol_name, chap_name, main_content)

            # 2. 如果成功生成了 AI 总结，将其隐式保存到 meta 并在后台落盘
            if ai_summary:
                chap_data["ai_synopsis"] = ai_summary
                self.project.save_meta()
                self.memory.upsert_analysis(vol_name, chap_name, {"summary": ai_summary})
                self.refresh_memory_dashboard()

            # 3. 如果用户还停留在这个章节，确保文本框里显示的是纯净的、没有尾巴的正文
            if self.current_vol_index == self.gen_v_idx and self.current_chap_index == self.gen_c_idx:
                self.content_output.setPlainText(main_content)

        # 清除后台生成标记
        self.is_generating = False
        self.gen_v_idx = -1
        self.gen_c_idx = -1

        # 刷新 UI 状态恢复原貌
        self.update_ui_state()
        self.statusBar().showMessage("✅ 章节正文生成完毕，AI内部线索梗概已入库！", 3000)

    #追加:自动挂机类函数
    def toggle_auto_pilot(self, mode="full"):
        if mode == "stop":
            if getattr(self, 'is_generating_summaries', False):
                if hasattr(self, 'summary_worker') and self.summary_worker.isRunning():
                    self.summary_worker.cancel()
                self.btn_auto_pilot.setText("🛑 正在停止补全...")
                self.btn_auto_pilot.setEnabled(False)
                return
            if getattr(self, 'is_auto_piloting', False):
                if hasattr(self, 'auto_worker') and self.auto_worker.isRunning():
                    self.auto_worker.cancel()
                    self.btn_auto_pilot.setText("🛑 正在停止挂机...")
                    self.btn_auto_pilot.setEnabled(False)
                else:
                    self.auto_pilot_finished()
            return

        api_key = self.settings.value("api_key", "")
        if not api_key:
            QMessageBox.warning(self, "错误", "缺少 API Key！")
            return

        # 一键生成本卷的前置校验
        if mode == "volume":
            if self.current_vol_index == -1:
                QMessageBox.warning(self, "错误", "请先在左侧选择需要生成的一卷（或卷下的某章）！")
                return
            vol_syn = self.project.meta["volumes"][self.current_vol_index].get("synopsis", "").strip()
            if not vol_syn:
                QMessageBox.warning(self, "错误",
                                    "该卷梗概为空！\n请先在中间面板填写【本卷的核心主线、剧情走向梗概】，以便AI有据可依。")
                return

        msg = '确定开启全自动挂机？\nAI将自动消耗大量Token补全所有设定和正文！' if mode == "full" else '确定一键生成本卷？\nAI将基于当前卷梗概，自动为您扩展章节并撰写本卷全部正文！'
        reply = QMessageBox.question(self, '高能预警', msg,
                                     QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        if reply != QMessageBox.StandardButton.Yes: return

        self.save_all()

        # 补全前置总结的范围划分：
        # 如果是全书挂机，检查截止当前全书所有章节的概要缺失；
        # 如果是单卷挂机，只检查在选中卷之前发生的所有剧情总结。
        target_v = self.current_vol_index if mode == "volume" else None
        target_c = 0 if mode == "volume" else None

        self._check_and_fill_summaries(target_v, target_c, lambda: self._execute_auto_pilot(mode, target_v))

    def _execute_auto_pilot(self, mode, target_v_idx):
        self.is_auto_piloting = True
        self.update_ui_state()

        base_url = self.settings.value("base_url", "https://api.deepseek.com")
        ai_model = self.settings.value("model", "deepseek-reasoner")
        temp = float(self.settings.value("temperature", 0.7))

        self.auto_worker = AutoPilotWorker(
            self.settings.value("api_key", ""), base_url, ai_model, temp,
            self.project, mode=mode, target_v_idx=target_v_idx if target_v_idx is not None else -1
        )

        self.auto_worker.status_signal.connect(lambda msg: self.statusBar().showMessage(msg))
        self.auto_worker.log_signal.connect(lambda msg: self.thinking_output.append(msg))

        self.hit_summary_delimiter = False
        self.auto_worker.content_signal.connect(self.append_content)
        self.auto_worker.reasoning_signal.connect(self.append_thinking)
        self.auto_worker.start_chapter_signal.connect(self.auto_start_chapter, Qt.ConnectionType.BlockingQueuedConnection)
        self.auto_worker.add_volume_signal.connect(self.auto_add_volume, Qt.ConnectionType.BlockingQueuedConnection)
        self.auto_worker.add_chapter_signal.connect(self.auto_add_chapter, Qt.ConnectionType.BlockingQueuedConnection)
        self.auto_worker.save_content_signal.connect(self.auto_save_content, Qt.ConnectionType.BlockingQueuedConnection)
        self.auto_worker.update_chapter_signal.connect(self.auto_update_chapter, Qt.ConnectionType.BlockingQueuedConnection)
        self.auto_worker.update_volume_signal.connect(self.auto_update_volume, Qt.ConnectionType.BlockingQueuedConnection)

        self.auto_worker.finished_signal.connect(self.auto_pilot_finished)
        self.auto_worker.error_signal.connect(self.handle_error)

        self.auto_worker.start()

    def auto_update_volume(self, v_idx, synopsis):
        vol = self.project.meta["volumes"][v_idx]
        vol["synopsis"] = synopsis
        self.project.save_meta()

        # 如果当前 UI 正好停留在这一卷的设置界面，实时刷新文本框
        if self.current_vol_index == v_idx and self.stacked_widget.currentIndex() == 1:
            self.vol_synopsis_input.setText(synopsis)

    # --- 供 AutoPilotWorker 跨线程调用的 UI 和数据更新槽函数 ---
    def auto_update_chapter(self, v_idx, c_idx, ai_synopsis):
        chap = self.project.meta["volumes"][v_idx]["chapters"][c_idx]
        chap["ai_synopsis"] = ai_synopsis

        # 核心逻辑：如果用户原本就没有写 synopsis，那就把 AI 写的塞到台面上；
        # 如果用户写了，那就保留用户写的，AI 的扩写只放在隐式的 ai_synopsis 里供大模型看
        if not chap.get("synopsis", "").strip():
            chap["synopsis"] = ai_synopsis

        self.project.save_meta()

        # 如果当前 UI 正好停留在这一章，刷新一下文本框显示
        if self.current_vol_index == v_idx and self.current_chap_index == c_idx:
            self.chap_synopsis_input.setText(chap.get("synopsis", ""))

    def auto_update_events(self, v_idx, events):
        """后台收到大事件数据更新时，静默落盘保存到 meta"""
        self.project.meta["volumes"][v_idx]["events"] = events
        self.project.save_meta()

    def auto_add_volume(self, name, synopsis):
        self.project.add_volume(name, synopsis)
        self.refresh_tree()
        self.tree.scrollToBottom()

    def auto_add_chapter(self, v_idx, name, ai_synopsis):
        # 【修改处】将 ai_synopsis 同时也赋值给 synopsis 字段，这样就能在 UI 的“章设定”里看到了！
        self.project.add_chapter(v_idx, name, synopsis=ai_synopsis, ai_synopsis=ai_synopsis)
        self.refresh_tree()
        self.tree.scrollToBottom()

    def auto_start_chapter(self, v_idx, c_idx):
        self.gen_v_idx = v_idx
        self.gen_c_idx = c_idx
        self.gen_content_buffer = ""
        self.gen_reasoning_buffer = ""
        self.hit_summary_delimiter = False

        # 自动选中左侧树状图对应的章节节点
        root = self.tree.topLevelItem(0)
        if root and v_idx < root.childCount():
            v_node = root.child(v_idx)
            if c_idx < v_node.childCount():
                c_node = v_node.child(c_idx)
                # 选中树节点
                self.tree.setCurrentItem(c_node)
                # 触发点击事件，让右侧面板切换到该章的空白编辑状态
                self.on_tree_select(c_node)

    def auto_save_content(self, v_idx, c_idx, main_content, ai_summary):
        vol_name = self.project.meta["volumes"][v_idx]["name"]
        chap_name = self.project.meta["volumes"][v_idx]["chapters"][c_idx]["name"]

        # 保存本地 docx
        old_content = self.project.read_chapter_content(vol_name, chap_name)
        if old_content.strip() and old_content != main_content:
            self.memory.archive_chapter_version(vol_name, chap_name, old_content, "auto-pilot")
        self.project.save_chapter_content(vol_name, chap_name, main_content)
        # 更新 meta 中的 AI 总结
        if ai_summary:
            self.project.meta["volumes"][v_idx]["chapters"][c_idx]["ai_synopsis"] = ai_summary
            self.project.save_meta()
            self.memory.upsert_analysis(vol_name, chap_name, {"summary": ai_summary})
            self.refresh_memory_dashboard()

        # 【关键修复】：取消这行 clear()，将清理工作交给 on_tree_select 去自然过渡
        self.hit_summary_delimiter = False

    def auto_pilot_finished(self):
        self.is_auto_piloting = False
        self.update_ui_state()

    def cancel_correction(self):
        """手动暂停/终止纠错任务"""
        if getattr(self, 'is_correcting', False):
            if hasattr(self, 'correct_worker') and self.correct_worker.isRunning():
                self.correct_worker.cancel()  # 触发 Worker 内的取消标记，并强行切断网络流
                self.log_list.addItem("⚠️ 接收到停止指令，正在等待当前请求安全中断...")
                self.log_list.scrollToBottom()
                self.statusBar().showMessage("🛑 正在停止纠错...", 3000)

    # === 正文区右键菜单 & 侧边栏切换逻辑 ===
    def toggle_right_sidebar(self, page_index, clicked_btn):
        # 实现类似 VSCode 左侧栏的点击展开/折叠效果
        if self.sidebar_stacked.isVisible() and self.sidebar_stacked.currentIndex() == page_index:
            self.sidebar_stacked.hide()
            clicked_btn.setChecked(False)
        else:
            self.sidebar_stacked.setCurrentIndex(page_index)
            self.sidebar_stacked.show()
            self.btn_sidebar_log.setChecked(page_index == 0)
            self.btn_sidebar_modifier.setChecked(page_index == 1)
            if hasattr(self, "btn_sidebar_memory"):
                self.btn_sidebar_memory.setChecked(page_index == 2)

    def show_editor_context_menu(self, pos):
        # 调用 PyQt 原生的富文本标准菜单
        menu = self.content_output.createStandardContextMenu()
        cursor = self.content_output.textCursor()

        # 如果用户选中了文本，则在菜单最下方动态加上“文段修正”
        if cursor.hasSelection():
            menu.addSeparator()
            action_modify = menu.addAction("🪄 文段修正")
            # QAction 的触发连接
            action_modify.triggered.connect(self.open_modifier_for_selection)

        menu.exec(self.content_output.mapToGlobal(pos))

    def open_modifier_for_selection(self):
        # 1. 保存当前的游标位置，便于之后回填
        self.target_modify_cursor = self.content_output.textCursor()
        selected_text = self.target_modify_cursor.selectedText().replace('\u2029', '\n')

        # 2. 展开侧边面板
        self.sidebar_stacked.setCurrentIndex(1)
        self.sidebar_stacked.show()
        self.btn_sidebar_modifier.setChecked(True)
        self.btn_sidebar_log.setChecked(False)
        if hasattr(self, "btn_sidebar_memory"):
            self.btn_sidebar_memory.setChecked(False)

        # 3. 数据灌入
        self.mod_selected_text.setPlainText(selected_text)
        self.mod_instruction.clear()
        self.mod_result.clear()
        self.mod_instruction.setFocus()
        self.btn_apply_replace.setEnabled(False)

    # === AI 文段修正核心逻辑 ===
    def start_segment_modification(self):
        selected_text = self.mod_selected_text.toPlainText().strip()
        instruction = self.mod_instruction.toPlainText().strip()
        if not selected_text or not instruction:
            QMessageBox.warning(self, "提示", "待修改片段和修改指令都不能为空！")
            return

        api_key = self.settings.value("api_key", "")
        if not api_key:
            QMessageBox.warning(self, "错误", "缺少 API Key！")
            return

        # 提取全局和局部上下文
        global_story = self.project.meta.get("global_synopsis", "")
        char_texts = [f"【{c['name']}】 性别:{c['gender']} 性格:{c['personality']} 经历:{c['experience']}" for c in
                      self.project.meta.get("characters", [])]
        char_setting = "\n".join(char_texts) if char_texts else "未提供"
        full_context = self.content_output.toPlainText()

        sys_prompt = f"你是一位精通网文写作的顶级大神。请根据用户的指令，对给定的小说片段进行重写、扩写或润色。\n\n【全局故事大纲】\n{global_story}\n\n【核心人物设定】\n{char_setting}\n\n【要求】：只返回修改后的纯正文文本，禁止输出任何解释性的废话！"
        user_prompt = f"【本章完整上下文参考】\n{full_context}\n\n【待修改的目标文段】\n{selected_text}\n\n【使用者的修改指令】\n{instruction}\n\n请直接输出修改后的文本："

        # 界面状态切换
        self.mod_result.clear()
        self.btn_submit_modify.setEnabled(False)
        self.btn_submit_modify.setText("正在生成...")
        self.btn_cancel_modify.setEnabled(True)
        self.btn_apply_replace.setEnabled(False)

        base_url = self.settings.value("base_url", "https://api.deepseek.com")
        model = self.settings.value("model", "deepseek-reasoner")
        temp = float(self.settings.value("temperature", 0.7))

        self.mod_worker = SegmentModifyWorker(api_key, base_url, model, temp, sys_prompt, user_prompt)
        # 如果模型吐出了思考过程，我们可以拼接到原先的思考日志窗，或者直接无视
        self.mod_worker.reasoning_signal.connect(self.append_thinking)
        self.mod_worker.content_signal.connect(lambda text: self.mod_result.insertPlainText(text))
        self.mod_worker.finished_signal.connect(self.finish_segment_modification)
        self.mod_worker.error_signal.connect(lambda e: QMessageBox.critical(self, "错误", str(e)))

        self.mod_worker.start()

    def cancel_segment_modification(self):
        if hasattr(self, 'mod_worker') and self.mod_worker.isRunning():
            self.mod_worker.cancel()
        self.finish_segment_modification()

    def finish_segment_modification(self):
        self.btn_submit_modify.setEnabled(True)
        self.btn_submit_modify.setText("✨ 生成修改")
        self.btn_cancel_modify.setEnabled(False)
        # 只要有一点结果，就可以允许替换
        if self.mod_result.toPlainText().strip():
            self.btn_apply_replace.setEnabled(True)

    def apply_modification(self):
        new_text = self.mod_result.toPlainText().strip()
        if not new_text or not hasattr(self, 'target_modify_cursor'):
            return

        # 重新选中文本并直接覆盖替换
        self.target_modify_cursor.insertText(new_text)
        self.statusBar().showMessage("✅ 文段已成功替换，别忘了按 Ctrl+S 保存！", 3000)

        # 自动收起侧边栏
        self.sidebar_stacked.hide()
        self.btn_sidebar_modifier.setChecked(False)
        self.content_output.setFocus()
