# data_manager.py
import os
import json
import shutil
import docx

class NovelProject:
    def __init__(self, root_path):
        self.root_path = root_path
        self.meta_path = os.path.join(self.root_path, "meta.json")
        self.meta = {
            "title": os.path.basename(self.root_path),
            "global_synopsis": "",
            "characters": [],
            "volumes": []
        }
        self.load_meta()

    def load_meta(self):
        if os.path.exists(self.meta_path):
            with open(self.meta_path, 'r', encoding='utf-8') as f:
                self.meta = json.load(f)
        else:
            self.save_meta()

    def save_meta(self):
        with open(self.meta_path, 'w', encoding='utf-8') as f:
            json.dump(self.meta, f, ensure_ascii=False, indent=4)

    def add_volume(self, vol_name, synopsis=""):
        vol_path = os.path.join(self.root_path, vol_name)
        if not os.path.exists(vol_path):
            os.makedirs(vol_path)
        self.meta["volumes"].append({"name": vol_name, "synopsis": synopsis, "chapters": []})
        self.save_meta()

    def add_chapter(self, vol_index, chap_name, synopsis="", ai_synopsis=""):
        vol_name = self.meta["volumes"][vol_index]["name"]
        chap_filename = f"{chap_name}.docx"
        chap_path = os.path.join(self.root_path, vol_name, chap_filename)

        if not os.path.exists(chap_path):
            doc = docx.Document()
            doc.save(chap_path)

        self.meta["volumes"][vol_index]["chapters"].append({
            "name": chap_name,
            "synopsis": synopsis,
            "ai_synopsis": ai_synopsis
        })
        self.save_meta()

    def rename_volume(self, v_idx, new_name):
        old_name = self.meta["volumes"][v_idx]["name"]
        if old_name == new_name: return
        old_path = os.path.join(self.root_path, old_name)
        new_path = os.path.join(self.root_path, new_name)
        if os.path.exists(old_path):
            os.rename(old_path, new_path)
        self.meta["volumes"][v_idx]["name"] = new_name
        self.save_meta()

    def rename_chapter(self, v_idx, c_idx, new_name):
        vol_name = self.meta["volumes"][v_idx]["name"]
        old_name = self.meta["volumes"][v_idx]["chapters"][c_idx]["name"]
        if old_name == new_name: return
        old_path = os.path.join(self.root_path, vol_name, f"{old_name}.docx")
        new_path = os.path.join(self.root_path, vol_name, f"{new_name}.docx")
        if os.path.exists(old_path):
            os.rename(old_path, new_path)
        self.meta["volumes"][v_idx]["chapters"][c_idx]["name"] = new_name
        self.save_meta()

    def delete_volume(self, v_idx):
        if v_idx < 0 or v_idx >= len(self.meta["volumes"]):
            return None

        source_vol = self.meta["volumes"][v_idx]
        source_name = source_vol["name"]
        source_path = os.path.join(self.root_path, source_name)

        chapters_to_move = list(source_vol.get("chapters", []))
        target_idx = self._volume_delete_target_index(v_idx)

        # 如果这是唯一一个卷且里面有章节，先创建一个兜底卷承接章节。
        if target_idx is None and chapters_to_move:
            fallback_name = self._unique_volume_name("未分卷章节")
            fallback_path = os.path.join(self.root_path, fallback_name)
            os.makedirs(fallback_path, exist_ok=True)
            self.meta["volumes"].append({"name": fallback_name, "synopsis": "删除原卷时自动承接的章节。", "chapters": []})
            target_idx = len(self.meta["volumes"]) - 1

        if target_idx is not None and chapters_to_move:
            target_vol = self.meta["volumes"][target_idx]
            target_name = target_vol["name"]
            target_path = os.path.join(self.root_path, target_name)
            os.makedirs(target_path, exist_ok=True)

            for chap in chapters_to_move:
                old_chap_name = chap["name"]
                new_chap_name = self._unique_chapter_name_in_volume(target_idx, old_chap_name)
                old_path = os.path.join(source_path, f"{old_chap_name}.docx")
                new_path = os.path.join(target_path, f"{new_chap_name}.docx")
                if os.path.exists(old_path):
                    shutil.move(old_path, new_path)
                moved_chap = dict(chap)
                moved_chap["name"] = new_chap_name
                target_vol["chapters"].append(moved_chap)

        # 删除源卷元数据。如果刚创建兜底卷，它的位置可能在源卷之后；删除源卷后无需额外修正。
        del self.meta["volumes"][v_idx]

        if os.path.exists(source_path):
            shutil.rmtree(source_path)
        self.save_meta()
        return target_idx

    def _volume_delete_target_index(self, v_idx):
        if len(self.meta["volumes"]) <= 1:
            return None
        if v_idx > 0:
            return v_idx - 1
        return 1

    def _unique_volume_name(self, base_name):
        existing = {v["name"] for v in self.meta["volumes"]}
        if base_name not in existing:
            return base_name
        counter = 2
        while f"{base_name}-{counter}" in existing:
            counter += 1
        return f"{base_name}-{counter}"

    def _unique_chapter_name_in_volume(self, v_idx, base_name):
        existing = {c["name"] for c in self.meta["volumes"][v_idx].get("chapters", [])}
        if base_name not in existing:
            return base_name
        counter = 2
        while f"{base_name}-{counter}" in existing:
            counter += 1
        return f"{base_name}-{counter}"

    def delete_chapter(self, v_idx, c_idx):
        vol_name = self.meta["volumes"][v_idx]["name"]
        chap_name = self.meta["volumes"][v_idx]["chapters"][c_idx]["name"]
        chap_path = os.path.join(self.root_path, vol_name, f"{chap_name}.docx")
        if os.path.exists(chap_path):
            os.remove(chap_path)
        del self.meta["volumes"][v_idx]["chapters"][c_idx]
        self.save_meta()

    def read_chapter_content(self, vol_name, chap_name):
        chap_path = os.path.join(self.root_path, vol_name, f"{chap_name}.docx")
        if os.path.exists(chap_path):
            doc = docx.Document(chap_path)
            return "\n".join([p.text for p in doc.paragraphs])
        return ""

    def save_chapter_content(self, vol_name, chap_name, content):
        chap_path = os.path.join(self.root_path, vol_name, f"{chap_name}.docx")
        doc = docx.Document()
        for line in content.split('\n'):
            doc.add_paragraph(line)
        doc.save(chap_path)
