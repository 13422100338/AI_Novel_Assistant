"""Book export rendering and file writers."""

import html

def iter_project_chapters(project):
    """Yield ``(volume, chapter, content)`` from a NovelProject-like object."""
    for volume in project.meta["volumes"]:
        for chapter in volume["chapters"]:
            content = project.read_chapter_content(volume["name"], chapter["name"])
            yield volume, chapter, content


def render_txt(project, title: str) -> str:
    """Render the whole book as plain TXT content."""
    chunks = [f"《{title}》", ""]
    for volume in project.meta["volumes"]:
        chunks.extend([f"【{volume['name']}】", ""])
        for chapter in volume["chapters"]:
            chunks.extend([f"  {chapter['name']}", ""])
            content = project.read_chapter_content(volume["name"], chapter["name"])
            chunks.extend([content, ""])
    return "\n".join(chunks)


def render_markdown(project, title: str) -> str:
    """Render the whole book as Markdown content."""
    chunks = [f"# {title}", ""]
    for volume in project.meta["volumes"]:
        chunks.extend([f"## {volume['name']}", ""])
        for chapter in volume["chapters"]:
            content = project.read_chapter_content(volume["name"], chapter["name"])
            chunks.extend([f"### {chapter['name']}", "", content, ""])
    return "\n".join(chunks)


def render_pdf_html(project, title: str) -> str:
    """Render the whole book as simple HTML for Qt PDF printing."""
    html_content = f"<h1 style='text-align: center;'>{html.escape(title)}</h1>"
    for volume in project.meta["volumes"]:
        html_content += f"<h2 style='color: #2C3E50;'>{html.escape(volume['name'])}</h2>"
        for chapter in volume["chapters"]:
            html_content += f"<h3>{html.escape(chapter['name'])}</h3>"
            content = project.read_chapter_content(volume["name"], chapter["name"])
            for line in content.split("\n"):
                if line.strip():
                    html_content += (
                        "<p style='text-indent: 2em; line-height: 1.5;'>"
                        f"{html.escape(line.strip())}"
                        "</p>"
                    )
    return html_content


def export_docx(project, file_path: str, title: str):
    import docx

    document = docx.Document()
    document.add_heading(title, 0)

    for volume in project.meta["volumes"]:
        document.add_heading(volume["name"], level=1)
        for chapter in volume["chapters"]:
            document.add_heading(chapter["name"], level=2)
            content = project.read_chapter_content(volume["name"], chapter["name"])
            for line in content.split("\n"):
                if line.strip():
                    document.add_paragraph(line.strip())

    document.save(file_path)


def export_txt(project, file_path: str, title: str):
    with open(file_path, "w", encoding="utf-8") as f:
        f.write(render_txt(project, title))


def export_markdown(project, file_path: str, title: str):
    with open(file_path, "w", encoding="utf-8") as f:
        f.write(render_markdown(project, title))


def export_pdf(project, file_path: str, title: str):
    from PyQt6.QtGui import QTextDocument
    from PyQt6.QtPrintSupport import QPrinter

    document = QTextDocument()
    document.setHtml(render_pdf_html(project, title))

    printer = QPrinter(QPrinter.PrinterMode.HighResolution)
    printer.setOutputFormat(QPrinter.OutputFormat.PdfFormat)
    printer.setOutputFileName(file_path)

    document.print(printer)
